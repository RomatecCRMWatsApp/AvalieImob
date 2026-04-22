# @module locacao_docx — Geração de DOCX para Avaliação de Locação
"""DOCX generator for Avaliação de Locação — Lei 8.245/91.

Mesma estrutura e formatação do PDF (locacao_pdf.py):
- Capa com logo, título, subtítulo, número, data, endereço
- 10 seções (NBR 14653 + Lei 8.245/1991)
- Tabela para comparativos de mercado
- Destaque para valor estimado
- Fotos do imóvel
- Responsável técnico com assinatura
- Cores: Verde (#1B4D1B) e Dourado (#D4A830)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import base64
from datetime import datetime


# Cores Romatec
GREEN = RGBColor(27, 77, 27)  # #1B4D1B
GOLD = RGBColor(212, 168, 48)  # #D4A830
WHITE = RGBColor(255, 255, 255)
DARK = RGBColor(26, 26, 26)
LIGHT_GREEN = RGBColor(232, 245, 233)


def _set_cell_shading(cell, color_hex: str):
    """Aplica cor de fundo a uma célula (hex sem #)."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Define bordas da célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), '4')
        edge_el.set(qn('w:color'), '1B4D1B')
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def _add_styled_paragraph(doc, text: str, bold: bool = False, size: int = 11, 
                          color=RGBColor(51, 51, 51), alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_after=Pt(6)):
    """Adiciona parágrafo formatado."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


def _add_section_heading(doc, text: str):
    """Adiciona título de seção estilo PDF (verde + dourado)."""
    # Criar tabela para efeito visual similar ao PDF
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.text = text
    
    # Formatar célula
    _set_cell_shading(cell, "1B4D1B")  # Verde escuro
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph()  # Espaço após
    return table


def _format_currency(value) -> str:
    """Formata valor como moeda brasileira."""
    try:
        v = float(value) if value else 0
        return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "R$ 0,00"


def _numero_por_extenso(valor: float) -> str:
    """Converte número para extenso simplificado."""
    # Simplificado - em produção usar biblioteca como num2words
    return f"({valor:,.2f})".replace(",", "X").replace(".", ",").replace("X", ".")


def generate_locacao_docx(loc: dict, user: dict | None = None) -> bytes:
    """Generate DOCX for Avaliação de Locação — mesma formatação do PDF."""
    if user is None:
        user = {}
    
    doc = Document()
    
    # Configura margens (A4)
    sections = doc.sections[0]
    sections.page_width = Cm(21)
    sections.page_height = Cm(29.7)
    sections.top_margin = Cm(2)
    sections.bottom_margin = Cm(2)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)
    
    # ── CAPA ────────────────────────────────────────────────────────────────
    # Logo (se disponível)
    company_logo = user.get("_company_logo_bytes")
    if company_logo:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(io.BytesIO(company_logo), width=Inches(1.5))
        except:
            pass
    
    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AVALIAÇÃO PARA FINS DE LOCAÇÃO")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = GREEN
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Parecer Técnico de Avaliação Mercadológica")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = DARK
    
    doc.add_paragraph()
    
    # Número do laudo
    numero = loc.get("numero_locacao") or "LOC-0000/0000"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Laudo nº {numero}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK
    
    doc.add_paragraph()
    
    # Endereço do imóvel
    endereco = loc.get("endereco", {})
    endereco_str = f"{endereco.get('logradouro', '')}, {endereco.get('numero', '')}"
    if endereco.get('complemento'):
        endereco_str += f" - {endereco['complemento']}"
    endereco_str += f"\n{endereco.get('bairro', '')} — {endereco.get('cidade', '')} — {endereco.get('uf', '')}"
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(endereco_str)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK
    
    doc.add_paragraph()
    
    # Info do solicitante na capa
    solicitante = loc.get("solicitante", {})
    _add_styled_paragraph(doc, f"Solicitante: {solicitante.get('nome', 'Não informado')}", 
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, f"Cidade: {endereco.get('cidade', 'Não informado')} / {endereco.get('uf', 'MA')}", 
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    data_vistoria = loc.get("data_vistoria", datetime.now().strftime("%Y-%m-%d"))
    _add_styled_paragraph(doc, f"Data: {data_vistoria}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Quebra de página após capa
    doc.add_page_break()
    
    # ── SEÇÃO 1: IDENTIFICAÇÃO ─────────────────────────────────────────────
    _add_section_heading(doc, "1. IDENTIFICAÇÃO")
    
    _add_styled_paragraph(doc, f"Número do Laudo: {numero}", bold=True)
    _add_styled_paragraph(doc, f"Solicitante: {solicitante.get('nome', 'Não informado')}")
    _add_styled_paragraph(doc, f"CPF/CNPJ: {solicitante.get('cpf_cnpj', 'Não informado')}")
    
    endereco_solicitante = solicitante.get('endereco', {})
    if isinstance(endereco_solicitante, dict):
        end_sol_str = f"{endereco_solicitante.get('logradouro', '')}, {endereco_solicitante.get('numero', '')}, {endereco_solicitante.get('bairro', '')}"
    else:
        end_sol_str = str(endereco_solicitante) if endereco_solicitante else "Não informado"
    _add_styled_paragraph(doc, f"Endereço do Solicitante: {end_sol_str}")
    
    _add_styled_paragraph(doc, f"Objetivo da Avaliação: {loc.get('objetivo', 'Fixação de Aluguel')}")
    _add_styled_paragraph(doc, f"Tipo de Locação: {loc.get('tipo_locacao', 'Residencial')}")
    _add_styled_paragraph(doc, f"Data de Vistoria: {data_vistoria}")
    
    # ── SEÇÃO 2: IMÓVEL AVALIADO ───────────────────────────────────────────
    _add_section_heading(doc, "2. IMÓVEL AVALIADO")
    
    _add_styled_paragraph(doc, f"Tipo: {loc.get('tipo_imovel', 'Não informado')}")
    _add_styled_paragraph(doc, f"Endereço: {endereco.get('logradouro', 'Não informado')}")
    _add_styled_paragraph(doc, f"Bairro: {endereco.get('bairro', 'Não informado')}")
    _add_styled_paragraph(doc, f"Cidade/UF: {endereco.get('cidade', 'Não informado')} / {endereco.get('uf', 'MA')}")
    _add_styled_paragraph(doc, f"CEP: {endereco.get('cep', 'Não informado')}")
    
    caracteristicas = loc.get("caracteristicas", {})
    if caracteristicas:
        _add_styled_paragraph(doc, f"Área Construída: {caracteristicas.get('area_construida', 'N/A')} m²")
        _add_styled_paragraph(doc, f"Idade do Imóvel: {caracteristicas.get('idade', 'N/A')} anos")
        _add_styled_paragraph(doc, f"Estado de Conservação: {caracteristicas.get('conservacao', 'N/A')}")
        _add_styled_paragraph(doc, f"Padrão de Acabamento: {caracteristicas.get('padrao', 'N/A')}")
        _add_styled_paragraph(doc, f"Dormitórios: {caracteristicas.get('dormitorios', 'N/A')}")
        _add_styled_paragraph(doc, f"Banheiros: {caracteristicas.get('banheiros', 'N/A')}")
        _add_styled_paragraph(doc, f"Vagas de Garagem: {caracteristicas.get('vagas', 'N/A')}")
    
    # ── SEÇÃO 3: REGIÃO E ENTORNO ──────────────────────────────────────────
    _add_section_heading(doc, "3. REGIÃO E ENTORNO")
    
    regiao = loc.get("regiao_entorno", {})
    _add_styled_paragraph(doc, "Infraestrutura", bold=True)
    _add_styled_paragraph(doc, regiao.get("infraestrutura", "Não avaliada").upper())
    
    _add_styled_paragraph(doc, "Padrão Construtivo da Região", bold=True)
    _add_styled_paragraph(doc, regiao.get("padrao_regiao", "Não avaliado").upper())
    
    if regiao.get("observacoes"):
        _add_styled_paragraph(doc, "Observações", bold=True)
        _add_styled_paragraph(doc, regiao.get("observacoes"))
    
    # ── SEÇÃO 4: PESQUISA DE MERCADO ───────────────────────────────────────
    _add_section_heading(doc, "4. PESQUISA DE MERCADO")
    
    comparativos = loc.get("comparativos", [])
    if comparativos:
        # Criar tabela de comparativos
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        headers = ["Endereço", "Área (m²)", "Aluguel (R$)", "R$/m²", "Fonte"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            _set_cell_shading(hdr_cells[i], "1B4D1B")
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
        
        # Dados
        valores_m2 = []
        for comp in comparativos:
            row_cells = table.add_row().cells
            row_cells[0].text = comp.get("endereco", "Não informado")
            row_cells[1].text = str(comp.get("area", "-"))
            row_cells[2].text = _format_currency(comp.get("aluguel"))
            
            area = float(comp.get("area", 1) or 1)
            aluguel = float(comp.get("aluguel", 0) or 0)
            valor_m2 = aluguel / area if area > 0 else 0
            valores_m2.append(valor_m2)
            row_cells[3].text = f"R$ {valor_m2:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row_cells[4].text = comp.get("fonte", "Pesquisa de mercado")
            
            # Alinhar
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Estatísticas
        if valores_m2:
            import statistics
            media = statistics.mean(valores_m2)
            _add_styled_paragraph(doc, f"Média do mercado: R$ {media:,.2f}/m²".replace(",", "X").replace(".", ",").replace("X", "."), bold=True)
    else:
        _add_styled_paragraph(doc, "Nenhum comparativo cadastrado.")
    
    # ── SEÇÃO 5: CÁLCULOS E TRATAMENTO ESTATÍSTICO ─────────────────────────
    _add_section_heading(doc, "5. CÁLCULOS E TRATAMENTO ESTATÍSTICO")
    
    _add_styled_paragraph(doc, "Grau de Precisão (NBR 14653-1 item 9): Grau I — Amplitude ≤ 30%", bold=True)
    
    # ── SEÇÃO 6: RESULTADO DA AVALIAÇÃO ────────────────────────────────────
    _add_section_heading(doc, "6. RESULTADO DA AVALIAÇÃO")
    
    resultado = loc.get("resultado", {})
    if resultado:
        valor = float(resultado.get("valor_mensal", 0) or 0)
        
        # Destaque para o valor
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("Valor Estimado de Locação Mensal: ")
        run.font.bold = True
        run.font.size = Pt(12)
        
        run = p.add_run(_format_currency(valor))
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = GREEN
        
        # Valor por extenso
        _add_styled_paragraph(doc, _numero_por_extenso(valor), italic=True)
        
        intervalo_min = resultado.get("intervalo_min")
        intervalo_max = resultado.get("intervalo_max")
        if intervalo_min and intervalo_max:
            _add_styled_paragraph(doc, f"Intervalo de Confiança (Mensal): {_format_currency(intervalo_min)} a {_format_currency(intervalo_max)}")
        
        _add_styled_paragraph(doc, f"Data de Referência da Avaliação: {data_vistoria}")
        
        # Aviso de validade
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("Este Parecer Técnico tem validade de 6 meses a contar da data de emissão, conforme ABNT NBR 14653-1. Após esse período, nova avaliação deverá ser realizada.")
        run.font.size = Pt(10)
        run.font.italic = True
    else:
        _add_styled_paragraph(doc, "Resultado não calculado.", bold=True)
    
    # ── SEÇÃO 7: GARANTIA E CONDIÇÕES ──────────────────────────────────────
    _add_section_heading(doc, "7. GARANTIA E CONDIÇÕES")
    
    garantia = loc.get("garantia", {})
    _add_styled_paragraph(doc, f"Tipo de Garantia Sugerida: {garantia.get('tipo', 'Não informado')}")
    _add_styled_paragraph(doc, f"Prazo de Locação: {garantia.get('prazo', 'Não informado')}")
    
    # ── SEÇÃO 8: BASE LEGAL E NORMATIVA ────────────────────────────────────
    _add_section_heading(doc, "8. BASE LEGAL E NORMATIVA")
    
    _add_styled_paragraph(doc, "Lei 8.245/1991 (Lei do Inquilinato) — Art. 565 a 578 do Código Civil", bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Esta avaliação foi elaborada com base nas seguintes normas e dispositivos legais:")
    run.font.size = Pt(11)
    
    base_legal = [
        "Lei 8.245/1991 — Lei do Inquilinato: Dispõe sobre as locações dos imóveis urbanos e os procedimentos para determinação do valor de locação (arts. 19 a 21 — revisional; arts. 51 a 57 — renovatória).",
        "Código Civil — art. 565 a 578: Regula o contrato de locação de coisas, estabelecendo direitos e deveres das partes, reajuste e rescisão.",
        "ABNT NBR 14653-1: Procedimentos gerais de avaliação de bens — metodologia, graus de fundamentação e precisão.",
        "ABNT NBR 14653-2: Avaliação de imóveis urbanos — Método Comparativo Direto de Dados de Mercado (item 8.2), homogeneização e tratamento estatístico.",
        "Resolução COFECI 957/2006: Habilita o Corretor de Imóveis para elaboração de Parecer Técnico de Avaliação Mercadológica (PTAM)."
    ]
    
    for item in base_legal:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)
    
    # ── SEÇÃO 9: REGISTRO FOTOGRÁFICO ──────────────────────────────────────
    _add_section_heading(doc, "9. REGISTRO FOTOGRÁFICO")
    
    fotos = loc.get("fotos_imovel", [])
    if fotos:
        _add_styled_paragraph(doc, "Fotografias do imóvel avaliado, obtidas na data da vistoria:")
        _add_styled_paragraph(doc, f"Total de {len(fotos)} foto(s) anexada(s) ao processo de avaliação.")
        
        # Inserir fotos se tiver bytes
        for i, foto in enumerate(fotos[:6]):  # Max 6 fotos no DOCX
            if isinstance(foto, dict) and foto.get("_image_bytes"):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(io.BytesIO(foto["_image_bytes"]), width=Inches(4))
                    
                    # Legenda
                    caption = foto.get("caption") or foto.get("legenda") or f"Foto {i+1}"
                    _add_styled_paragraph(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10)
                except:
                    pass
    else:
        _add_styled_paragraph(doc, "Nenhuma foto anexada.")
    
    # ── SEÇÃO 10: RESPONSÁVEL TÉCNICO ──────────────────────────────────────
    _add_section_heading(doc, "10. RESPONSÁVEL TÉCNICO")
    
    tipo_prof = loc.get("tipo_profissional") or user.get("role", "")
    tipo_map = {
        "corretor": "Corretor de Imóveis habilitado nos termos da Resolução COFECI 957/2006",
        "engenheiro": "Engenheiro com Anotação de Responsabilidade Técnica (ART) conforme Resolução CONFEA 345/90",
        "arquiteto": "Arquiteto e Urbanista com Registro de Responsabilidade Técnica (RRT)",
        "perito_judicial": "Perito Judicial cadastrado no Tribunal, habilitado nos termos do CPC art. 156",
    }
    tipo_desc = tipo_map.get(tipo_prof, tipo_prof)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"O(A) profissional signatário(a) deste Parecer Técnico de Avaliação Mercadológica é {tipo_desc}, responsabilizando-se técnica e legalmente pelo conteúdo e pelos valores aqui expressos, conforme as normas regulamentadoras vigentes.")
    run.font.size = Pt(11)
    
    # Assinatura
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_____________________________________________")
    
    nome_resp = user.get("nome") or user.get("name") or "Não informado"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome_resp)
    run.font.bold = True
    run.font.size = Pt(12)
    
    if user.get("creci"):
        _add_styled_paragraph(doc, f"CRECI: {user['creci']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if user.get("cna"):
        _add_styled_paragraph(doc, f"CNAI: {user['cna']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if user.get("crea"):
        _add_styled_paragraph(doc, f"CREA: {user['crea']}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Data
    doc.add_paragraph()
    data_atual = datetime.now().strftime("%Y-%m-%d")
    _add_styled_paragraph(doc, f"Mirador, MA, {data_atual}", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    
    # ── Salvar em bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

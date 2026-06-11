# @module docx_gen.vistoria_averbacao_docx — DOCX da Vistoria de Obra para Averbação
"""
Espelha a estrutura do PDF (seções + 3 quadros) em Word, com parágrafos justificados.
Galeria embute as imagens quando os bytes são fornecidos em photo["_bytes"].
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from models.averbacao import ETAPAS_OBRA, DOCS_AVERBACAO, faixa_divergencia, calcular_averbacao
from services.vistoria_averbacao_relatorio import gerar_secoes_averbacao

GREEN = RGBColor(0x0C, 0x33, 0x20)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
SIT_LABEL = {"OK": "OK", "PEND": "PENDENTE", "NA": "N/A", "PENDENTE_AVALIACAO": "—"}


def _num(v, casas: int = 2) -> str:
    try:
        f = float(v)
    except (TypeError, ValueError):
        return "—"
    return f"{f:,.{casas}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _justify(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _heading(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(10)
    return p


def _set_cell(cell, texto, bold=False, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.size = Pt(9)
    if bg:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)


def _quadro_confronto(doc, av):
    conf = av.get("confronto") or {}
    div_pct = conf.get("divergencia_pct")
    faixa = faixa_divergencia(div_pct)
    bg = {"verde": "E8F5E9", "ambar": "FFF3E0", "vermelho": "FDECEC"}.get(faixa)
    doc.add_paragraph().add_run("Quadro de Confronto de Áreas").bold = True
    t = doc.add_table(rows=2, cols=5)
    t.style = "Table Grid"
    headers = ["Projeto", "Executada", "Matrícula", "Divergência m²", "Divergência %"]
    vals = [
        _num(conf.get("area_projeto_m2")), _num(conf.get("area_medida_m2")),
        _num(conf.get("area_matricula_m2")) if conf.get("area_matricula_m2") not in (None, "") else "—",
        _num(conf.get("divergencia_m2")) if conf.get("divergencia_m2") is not None else "—",
        (f"{_num(div_pct)}%" if div_pct is not None else "—"),
    ]
    for i, h in enumerate(headers):
        _set_cell(t.rows[0].cells[i], h, bold=True, bg="0C3320")
    for i, v in enumerate(vals):
        _set_cell(t.rows[1].cells[i], v, bold=(i == 4), bg=(bg if i == 4 else None))


def _quadro_etapas(doc, av):
    etapas = {e.get("etapa_id"): e.get("percentual", 0) for e in (av.get("etapas") or []) if isinstance(e, dict)}
    doc.add_paragraph().add_run("Quadro de Etapas (NBR 12721 simplificado)").bold = True
    t = doc.add_table(rows=1 + len(ETAPAS_OBRA) + 1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["Etapa", "Peso", "% Executado"]):
        _set_cell(t.rows[0].cells[i], h, bold=True, bg="0C3320")
    r = 1
    for e in ETAPAS_OBRA:
        _set_cell(t.rows[r].cells[0], e["nome"])
        _set_cell(t.rows[r].cells[1], str(e["peso"]))
        _set_cell(t.rows[r].cells[2], f"{int(etapas.get(e['id'], 0))}%")
        r += 1
    _set_cell(t.rows[r].cells[0], "CONCLUSÃO GERAL", bold=True, bg="C9A84C")
    _set_cell(t.rows[r].cells[1], "100", bold=True, bg="C9A84C")
    _set_cell(t.rows[r].cells[2], f"{_num(av.get('conclusao_geral_pct'), 1)}%", bold=True, bg="C9A84C")


def _quadro_documental(doc, av):
    so_res = (av.get("destinacao") or "residencial") == "residencial"
    docs = {d.get("doc_id"): d for d in (av.get("documentos") or []) if isinstance(d, dict)}
    visiveis = [d for d in DOCS_AVERBACAO if not (d.get("comercial") and so_res)]
    doc.add_paragraph().add_run("Quadro Documental para Averbação").bold = True
    t = doc.add_table(rows=1 + len(visiveis), cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["Documento", "Base legal", "Situação"]):
        _set_cell(t.rows[0].cells[i], h, bold=True, bg="0C3320")
    r = 1
    for d in visiveis:
        item = docs.get(d["id"]) or {}
        sit = item.get("situacao", "PENDENTE_AVALIACAO")
        bg = {"OK": "E8F5E9", "PEND": "FFF3E0", "NA": "F1F1F1"}.get(sit)
        _set_cell(t.rows[r].cells[0], d["nome"])
        _set_cell(t.rows[r].cells[1], d["base"])
        _set_cell(t.rows[r].cells[2], SIT_LABEL.get(sit, "—"), bold=True, bg=bg)
        r += 1


def generate_averbacao_docx(vistoria: dict, user: dict, photos: list | None = None,
                            signatures: list | None = None) -> bytes:
    av = calcular_averbacao(dict(vistoria.get("averbacao") or {}))
    vistoria = dict(vistoria)
    vistoria["averbacao"] = av

    doc = Document()
    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("VISTORIA TÉCNICA DE OBRA PARA AVERBAÇÃO")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = GREEN
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Nº {vistoria.get('numero_tvi') or '—'}  ·  {vistoria.get('imovel_endereco') or ''}").font.size = Pt(11)
    doc.add_paragraph()

    secoes = gerar_secoes_averbacao(vistoria)
    for i, (titulo, corpo) in enumerate(secoes, 1):
        _heading(doc, f"{i}. {titulo}")
        for par in (corpo or "").split("\n"):
            if par.strip():
                p = doc.add_paragraph(par.strip())
                _justify(p)
                for r in p.runs:
                    r.font.size = Pt(10)
        if titulo == "CONFRONTO DE ÁREAS":
            _quadro_confronto(doc, av)
        elif titulo == "ESTÁGIO DA OBRA":
            _quadro_etapas(doc, av)
        elif titulo == "DOCUMENTAÇÃO PARA AVERBAÇÃO":
            _quadro_documental(doc, av)

    # Galeria
    fotos = photos or []
    if fotos:
        _heading(doc, "REGISTRO FOTOGRÁFICO")
        for ph in fotos:
            raw = ph.get("_bytes") if isinstance(ph, dict) else None
            if raw:
                try:
                    doc.add_picture(io.BytesIO(raw), width=Cm(10))
                except Exception:
                    pass
            cap = (ph.get("legenda") or ph.get("ambiente") or "") if isinstance(ph, dict) else ""
            if cap:
                c = doc.add_paragraph(cap)
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in c.runs:
                    r.font.size = Pt(8)
                    r.italic = True

    # Assinatura
    doc.add_paragraph()
    data_str = vistoria.get("data_vistoria") or datetime.utcnow().strftime("%d/%m/%Y")
    cidade = vistoria.get("imovel_cidade") or ""
    loc = doc.add_paragraph(f"{cidade}, {data_str}." if cidade else data_str)
    loc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    nome = vistoria.get("responsavel_nome") or (user.get("name") or "")
    assin = doc.add_paragraph()
    assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    assin.add_run("__________________________________________\n").font.size = Pt(10)
    rn = assin.add_run(nome)
    rn.bold = True
    rn.font.size = Pt(10)
    if vistoria.get("art_trt_numero"):
        trt = doc.add_paragraph(f"ART/TRT nº {vistoria['art_trt_numero']}")
        trt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# @module docx.tvi_docx — Geração de DOCX para Termo de Vistoria de Imóvel (TVI)
"""DOCX generator for TVI — mesmo padrão visual do locacao_docx.py.

Layout:
- Capa com logo, título, número TVI, endereço, solicitante, responsável técnico
- Seções numeradas com header verde (tabela 1×1, fundo #1B4D1B, texto branco)
- Tabelas formatadas com cabeçalho verde
- Fotos em grade 2 colunas com legenda
- Bloco de assinatura com imagem, linha, nome, CREA/CAU/CFTMA, data/local
- Cores: Verde (#1B4D1B) e Dourado (#D4A830)
"""
from __future__ import annotations

import base64
import io
import urllib.request
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ── Cores Romatec ─────────────────────────────────────────────────────────────
GREEN = RGBColor(27, 77, 27)        # #1B4D1B
GOLD = RGBColor(212, 168, 48)       # #D4A830
WHITE = RGBColor(255, 255, 255)
DARK = RGBColor(26, 26, 26)
LIGHT_GREEN = RGBColor(232, 245, 233)

# aliases mantidos para compatibilidade interna
BRAND_GREEN = GREEN
BRAND_GOLD = GOLD


# ── helpers de célula ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex: str) -> None:
    """Aplica cor de fundo a uma célula (hex sem #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


# alias para manter compatibilidade com o nome usado em locacao_docx
_set_cell_shading = _set_cell_bg


def _set_cell_border(cell) -> None:
    """Define bordas verdes na célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "4")
        edge_el.set(qn("w:color"), "1B4D1B")
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


# ── helpers de parágrafo ──────────────────────────────────────────────────────

def _add_styled_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                           size: int = 11, color=RGBColor(51, 51, 51),
                           alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    """Adiciona parágrafo formatado — igual locacao_docx."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def _add_section_heading(doc, text: str):
    """Título de seção: tabela 1×1 com fundo verde escuro, texto branco — igual locacao_docx."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.text = text
    _set_cell_shading(cell, "1B4D1B")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()
    return table


def _fmt_val(val: Any) -> str:
    """Formata valor de campo para exibição em texto plano."""
    if val is None:
        return ""
    if isinstance(val, bool):
        return "Sim" if val else "Não"
    if isinstance(val, list):
        items = [str(i) for i in val if i is not None and str(i).strip()]
        return ", ".join(items)
    if isinstance(val, dict):
        parts = [f"{k}: {v}" for k, v in val.items() if v is not None and str(v).strip()]
        return " | ".join(parts)
    return str(val)


def _add_label_value(doc: Document, label: str, value: Any) -> None:
    """Adiciona linha label: valor."""
    text = _fmt_val(value)
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"


def _add_para(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"


# ── renderers de seção ────────────────────────────────────────────────────────

def _render_cover(doc: Document, v: dict, user: dict, model_nome: str) -> None:
    """Capa: logo, título, número TVI, endereço, solicitante, responsável técnico."""
    # Logo da empresa
    company_logo = user.get("_company_logo_bytes")
    if company_logo:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(io.BytesIO(company_logo), width=Inches(1.5))
        except Exception:
            pass

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TERMO DE VISTORIA DE IMÓVEL")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = GREEN

    # Nome do modelo (subtítulo)
    if model_nome:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(model_nome.upper())
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = DARK

    doc.add_paragraph()

    # Número TVI
    numero = v.get("numero_tvi") or "TVI-0000/0000"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Nº {numero}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK

    doc.add_paragraph()

    # Endereço do imóvel
    endereco = v.get("imovel_endereco") or ""
    bairro = v.get("imovel_bairro") or ""
    city_uf = " / ".join(filter(None, [v.get("imovel_cidade"), v.get("imovel_uf")]))
    if endereco:
        parts = [endereco]
        if bairro:
            parts.append(bairro)
        if city_uf:
            parts.append(city_uf)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n".join(parts))
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = DARK

    doc.add_paragraph()

    # Solicitante / data
    _add_styled_paragraph(doc, f"Solicitante: {v.get('cliente_nome', 'Não informado')}",
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if city_uf:
        _add_styled_paragraph(doc, f"Cidade: {city_uf}", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    data_ref = v.get("data_vistoria") or datetime.now().strftime("%d/%m/%Y")
    _add_styled_paragraph(doc, f"Data: {data_ref}", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()


def _render_identificacao(doc: Document, v: dict) -> None:
    """Seção 1 — Identificação."""
    _add_section_heading(doc, "1. IDENTIFICAÇÃO")

    _add_label_value(doc, "Número TVI", v.get("numero_tvi"))
    _add_label_value(doc, "Solicitante", v.get("cliente_nome"))
    _add_label_value(doc, "CPF / CNPJ", v.get("cliente_cpf_cnpj"))
    _add_label_value(doc, "Telefone", v.get("cliente_telefone"))
    _add_label_value(doc, "E-mail", v.get("cliente_email"))
    doc.add_paragraph()
    _add_label_value(doc, "Responsável Técnico", v.get("responsavel_nome"))
    _add_label_value(doc, "CREA", v.get("responsavel_crea"))
    _add_label_value(doc, "CAU", v.get("responsavel_cau"))
    _add_label_value(doc, "CFTMA", v.get("responsavel_cftma"))
    _add_label_value(doc, "ART / TRT nº", v.get("art_trt_numero"))


def _render_imovel(doc: Document, v: dict) -> None:
    """Seção 2 — Dados do Imóvel."""
    _add_section_heading(doc, "2. DADOS DO IMÓVEL")

    _add_label_value(doc, "Tipo", v.get("imovel_tipo"))
    _add_label_value(doc, "Endereço completo", v.get("imovel_endereco"))
    _add_label_value(doc, "Bairro", v.get("imovel_bairro"))
    city_uf = " / ".join(filter(None, [v.get("imovel_cidade"), v.get("imovel_uf")]))
    _add_label_value(doc, "Cidade / UF", city_uf)
    _add_label_value(doc, "CEP", v.get("imovel_cep"))
    _add_label_value(doc, "Matrícula", v.get("imovel_matricula"))


def _render_vistoria_dados(doc: Document, v: dict) -> None:
    """Seção 3 — Dados da Vistoria."""
    _add_section_heading(doc, "3. DADOS DA VISTORIA")

    _add_label_value(doc, "Data", v.get("data_vistoria"))
    _add_label_value(doc, "Hora", v.get("hora_vistoria"))
    _add_label_value(doc, "Condições Climáticas", v.get("condicoes_climaticas"))
    _add_label_value(doc, "Objetivo", v.get("objetivo"))
    _add_label_value(doc, "Metodologia", v.get("metodologia"))


def _render_ambientes(doc: Document, v: dict) -> None:
    """Seção 4 — Ambientes Vistoriados."""
    ambientes = v.get("ambientes") or []
    if not ambientes:
        return
    _add_section_heading(doc, "4. AMBIENTES VISTORIADOS")

    for amb in ambientes:
        if not (isinstance(amb, dict) and amb.get("nome")):
            continue
        # Sub-heading do ambiente
        p = doc.add_paragraph()
        run = p.add_run(amb["nome"])
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = GREEN

        _add_label_value(doc, "Descrição", amb.get("descricao"))
        _add_label_value(doc, "Estado de Conservação", amb.get("estado_conservacao"))
        _add_label_value(doc, "Observações", amb.get("observacoes"))
        doc.add_paragraph()


def _render_campos_extras(doc: Document, v: dict,
                           campos_especificos: list | None = None) -> None:
    """Seção 5 — Campos Específicos do Modelo."""
    extras = v.get("campos_extras") or {}
    if not extras:
        return

    # Mapa id→{label, secao} do modelo
    campo_meta: dict[str, dict] = {}
    if campos_especificos:
        for c in campos_especificos:
            cid = c.get("id") or c.get("key") or ""
            if cid:
                campo_meta[cid] = {
                    "label": c.get("label") or cid.replace("_", " ").title(),
                    "secao": c.get("secao") or "Informações Complementares",
                }

    # Agrupa por seção
    secoes: dict[str, list] = {}
    for key, val in extras.items():
        text = _fmt_val(val)
        if not text or not text.strip():
            continue
        meta = campo_meta.get(key, {})
        label = meta.get("label") or key.replace("_", " ").title()
        secao = meta.get("secao") or "Informações Complementares"
        secoes.setdefault(secao, []).append((label, text))

    if not secoes:
        return

    _add_section_heading(doc, "5. CAMPOS ESPECÍFICOS DO MODELO")
    for secao, itens in secoes.items():
        p = doc.add_paragraph(secao)
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = GREEN
        for label, text in itens:
            _add_label_value(doc, label, text)


def _render_fotos(doc: Document, photos: list) -> None:
    """Seção 6 — Registro Fotográfico (grade 2 colunas, cabeçalho verde)."""
    if not photos:
        return
    _add_section_heading(doc, "6. REGISTRO FOTOGRÁFICO")
    _add_styled_paragraph(doc,
                          f"Fotografias obtidas na data da vistoria. Total: {len(photos)} foto(s).")

    MAX_W = Cm(7.5)
    MAX_H = Cm(5.5)

    # Agrupa em pares para tabela 2 colunas
    row_pairs: list = []
    pair: list = []
    for photo in photos:
        url = photo.get("url", "") if isinstance(photo, dict) else ""
        legenda = photo.get("legenda", "") if isinstance(photo, dict) else ""
        ambiente = photo.get("ambiente", "") if isinstance(photo, dict) else ""
        caption = legenda or ambiente or ""

        img_bytes: bytes | None = None
        if url:
            try:
                if url.startswith("data:"):
                    img_bytes = base64.b64decode(url.split(",", 1)[-1])
                else:
                    with urllib.request.urlopen(url, timeout=5) as resp:
                        img_bytes = resp.read()
            except Exception:
                img_bytes = None

        # também aceita _image_bytes direto
        if not img_bytes and isinstance(photo, dict) and photo.get("_image_bytes"):
            img_bytes = photo["_image_bytes"]

        pair.append({"img": img_bytes, "caption": caption})
        if len(pair) == 2:
            row_pairs.append(pair)
            pair = []
    if pair:
        while len(pair) < 2:
            pair.append({"img": None, "caption": ""})
        row_pairs.append(pair)

    for pair in row_pairs:
        tbl = doc.add_table(rows=2, cols=2)
        tbl.style = "Table Grid"
        for col_idx, cell_data in enumerate(pair):
            img_cell = tbl.rows[0].cells[col_idx]
            img_cell.width = Cm(8.5)
            cap_cell = tbl.rows[1].cells[col_idx]

            if cell_data["img"]:
                try:
                    para = img_cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(io.BytesIO(cell_data["img"]), width=MAX_W, height=MAX_H)
                except Exception:
                    img_cell.text = "[Foto indisponível]"

            cap_para = cap_cell.paragraphs[0]
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(cell_data["caption"])
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.name = "Calibri"

        doc.add_paragraph()


def _render_conclusao(doc: Document, v: dict) -> None:
    """Seção 7 — Conclusão Técnica."""
    _add_section_heading(doc, "7. CONCLUSÃO TÉCNICA")
    if v.get("conclusao_tecnica"):
        _add_para(doc, v["conclusao_tecnica"])
    else:
        _add_styled_paragraph(doc, "Conclusão não informada.")

    # Aviso de validade
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Este Termo de Vistoria tem validade de 90 dias a contar da data de emissão. "
        "Após esse período, nova vistoria deverá ser realizada para reavaliação das condições do imóvel."
    )
    run.font.size = Pt(10)
    run.font.italic = True


def _render_assinatura(doc: Document, v: dict, sigs: list) -> None:
    """Bloco de assinatura com local, data, imagem de assinatura, linha, nome, registros."""
    doc.add_paragraph()
    doc.add_paragraph()

    city = v.get("imovel_cidade", "")
    date_str = v.get("data_vistoria") or datetime.now().strftime("%d/%m/%Y")
    loc_text = f"{city}, {date_str}." if city else date_str
    loc_para = doc.add_paragraph(loc_text)
    loc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    # Imagem de assinatura (se disponível)
    first_sig = next(
        (s for s in (sigs or []) if isinstance(s, dict) and s.get("data_b64")), None
    )
    if first_sig:
        try:
            sig_data = base64.b64decode(first_sig["data_b64"])
            sig_para = doc.add_paragraph()
            sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sig_para.add_run()
            run.add_picture(io.BytesIO(sig_data), width=Cm(5.5), height=Cm(2.0))
        except Exception:
            pass

    sig_line = doc.add_paragraph("_" * 50)
    sig_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nome = v.get("responsavel_nome") or (first_sig or {}).get("signatario", "")
    if nome:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(nome)
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = "Calibri"

    for field, label in [
        ("responsavel_crea", "CREA"),
        ("responsavel_cau", "CAU"),
        ("responsavel_cftma", "CFTMA"),
        ("art_trt_numero", "ART/TRT nº"),
    ]:
        val = v.get(field, "")
        if val:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}: {val}")
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    if first_sig and first_sig.get("cargo"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(first_sig["cargo"])
        run.font.size = Pt(10)
        run.font.name = "Calibri"


# ── entry point ───────────────────────────────────────────────────────────────

def generate_tvi_docx(
    vistoria: dict,
    user: dict,
    photos: list | None = None,
    signatures: list | None = None,
    model_nome: str = "",
    campos_especificos: list | None = None,
) -> bytes:
    """Gera DOCX do TVI com padrão visual locacao_docx. Retorna bytes.

    Args:
        vistoria: Documento de vistoria (dict).
        user: Usuário logado (dict).
        photos: Lista de fotos [{url, ambiente, legenda, _image_bytes}, …].
        signatures: Lista de assinaturas [{data_b64, signatario, cargo}, …].
        model_nome: Nome legível do modelo de vistoria.
        campos_especificos: Campos do modelo TVI [{id, label, secao}, …].
    """
    if user is None:
        user = {}

    doc = Document()

    # Margens A4
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # ── Capa ──────────────────────────────────────────────────────────────────
    _render_cover(doc, vistoria, user, model_nome)

    # ── Sumário ───────────────────────────────────────────────────────────────
    _add_section_heading(doc, "SUMÁRIO")
    _toc_items = [
        ("1",    "Identificação — Solicitante e Objetivo"),
        ("2",    "Identificação do Imóvel"),
        ("3",    "Dados da Vistoria"),
        ("4",    "Ambientes Vistoriados"),
        ("5",    "Campos Específicos / Itens Adicionais"),
        ("6",    "Conclusão e Observações"),
        ("A.I",  "Anexo I — Registro Fotográfico"),
        ("A.II", "Anexo II — Assinatura das Partes"),
    ]
    _tbl_toc = doc.add_table(rows=1, cols=2)
    _tbl_toc.style = "Table Grid"
    _hdr = _tbl_toc.rows[0].cells
    for _i, _h in enumerate(["Seção", "Título"]):
        _hdr[_i].text = _h
        _set_cell_shading(_hdr[_i], "1B4D1B")
        _hdr[_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _run in _hdr[_i].paragraphs[0].runs:
            _run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _run.font.bold = True
            _run.font.size = Pt(10)
    for _num, _title in _toc_items:
        _row = _tbl_toc.add_row().cells
        _row[0].text = _num
        _row[1].text = _title
        _row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── Seção 1: Identificação ─────────────────────────────────────────────
    _render_identificacao(doc, vistoria)
    doc.add_paragraph()

    # ── Seção 2: Imóvel ────────────────────────────────────────────────────
    _render_imovel(doc, vistoria)
    doc.add_paragraph()

    # ── Seção 3: Dados da Vistoria ─────────────────────────────────────────
    _render_vistoria_dados(doc, vistoria)
    doc.add_paragraph()

    # ── Seção 4: Ambientes ─────────────────────────────────────────────────
    _render_ambientes(doc, vistoria)
    doc.add_paragraph()

    # ── Seção 5: Campos Extras ─────────────────────────────────────────────
    _render_campos_extras(doc, vistoria, campos_especificos or [])
    doc.add_paragraph()

    # ── Seção 6: Fotos ─────────────────────────────────────────────────────
    _render_fotos(doc, photos or [])
    doc.add_paragraph()

    # ── Seção 7: Conclusão ─────────────────────────────────────────────────
    _render_conclusao(doc, vistoria)

    doc.add_page_break()

    # ── Assinatura ─────────────────────────────────────────────────────────
    _render_assinatura(doc, vistoria, signatures or [])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

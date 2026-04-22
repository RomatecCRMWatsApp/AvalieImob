# @module ptam_docx — Geração de DOCX para PTAM (Parecer Técnico de Avaliação Mercadológica)
"""DOCX generator for PTAM — padrão visual RomaTec, paridade total com ptam_pdf.py.

Seções:
  1. Identificação e Objetivo (solicitante completo, finalidade, dados judiciais)
  2. Documentação Analisada
  3. Identificação do Imóvel (endereço, matrículas, proprietários, áreas)
  4. Caracterização do Imóvel (área construída, conservação, cômodos)
  5. Vistoria Técnica
  6. Análise da Região
  7. Homogeneização e Amostras de Mercado
  8. Metodologia
  9. Ponderância (se aplicável)
 10. Método de Avaliação / Depreciação (se aplicável)
 11. Resultado da Avaliação (valor, intervalo, prazo, campo de arbítrio)
 12. Considerações, Ressalvas e Pressupostos
 13. Base Legal e Normativa
 14. Declaração de Responsabilidade Técnica + Assinatura
 15. Registro Fotográfico (fotos grandes, 2/página)
 16. Documentos Anexos
 17. Certidões das Partes (CND)
"""
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


# ── Cores Romatec ─────────────────────────────────────────────────────────────
GREEN = RGBColor(27, 77, 27)       # #1B4D1B
GOLD = RGBColor(212, 168, 48)      # #D4A830
WHITE = RGBColor(255, 255, 255)
DARK = RGBColor(26, 26, 26)
LIGHT_GREEN = RGBColor(232, 245, 233)


# ── helpers de célula ─────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "4")
        edge_el.set(qn("w:color"), "1B4D1B")
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


# ── helpers de parágrafo ──────────────────────────────────────────────────────

def _add_styled_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                           size: int = 11, color=RGBColor(51, 51, 51),
                           alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def _add_section_heading(doc, text: str):
    """Título de seção: tabela 1×1 com fundo verde escuro, texto branco."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.text = text
    _set_cell_shading(cell, "1B4D1B")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()
    return table


def _add_subsection_heading(doc, text: str):
    """Sub-título de seção em verde."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_label_value(doc, label: str, value, bold_label: bool = True, skip_zero: bool = True):
    """Adiciona linha 'label: valor'. Ignora vazios e zeros."""
    text = _fmt_val(value)
    if not text or not text.strip():
        return
    if skip_zero and text.strip() in ("0", "0.0", "0.00", "0,00", "False"):
        return
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = bold_label
    r1.font.size = Pt(11)
    r2 = p.add_run(str(text))
    r2.font.size = Pt(11)


def _fmt_val(val) -> str:
    if val is None:
        return ""
    if isinstance(val, bool):
        return "Sim" if val else ""
    if isinstance(val, list):
        return ", ".join(str(i) for i in val if i is not None and str(i).strip())
    if isinstance(val, dict):
        return " | ".join(f"{k}: {v}" for k, v in val.items() if v is not None and str(v).strip())
    return str(val)


def _format_currency(value) -> str:
    try:
        v = float(value) if value else 0
        return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "R$ 0,00"


def _format_area(value) -> str:
    try:
        return f"{float(value):,.2f} m²".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "0,00 m²"


def _format_area_ha(value) -> str:
    try:
        return f"{float(value):,.4f} ha".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "0,0000 ha"


# ── mapa de finalidade (igual ao PDF) ─────────────────────────────────────────
_FINALIDADE_MAP = {
    "cv_alienacao": "Compra e Venda — Alienação",
    "cv_aquisicao": "Compra e Venda — Aquisição",
    "cv_oferta": "Compra e Venda — Oferta Pública",
    "cv_dacao": "Compra e Venda — Dação em Pagamento",
    "gar_sfh": "Garantia Bancária — Financiamento SFH",
    "gar_sfi": "Garantia Bancária — Financiamento SFI",
    "gar_credito_rural": "Garantia Bancária — Crédito Rural (Penhor Rural)",
    "gar_refinanciamento": "Garantia Bancária — Refinanciamento",
    "gar_lci_cri": "Garantia Bancária — LCI / CRI",
    "gar_ccb": "Garantia Bancária — CCB Imobiliária",
    "judicial_partilha": "Judicial — Partilha de Bens (Inventário / Divórcio)",
    "judicial_desapropriacao": "Judicial — Desapropriação",
    "judicial_indenizacao": "Judicial — Ação de Indenização",
    "judicial_execucao": "Judicial — Execução de Sentença",
    "judicial_usucapiao": "Judicial — Usucapião",
    "judicial_pericia": "Perícia Judicial (CPC art. 156)",
    "loc_fixacao": "Locação — Fixação de Aluguel",
    "loc_revisao": "Locação — Revisão de Aluguel (Lei 8.245/91)",
    "loc_renovatoria": "Locação — Ação Renovatória",
    "seg_reposicao": "Seguros — Valor de Reposição",
    "seg_sinistro": "Seguros — Sinistro",
    "seg_risco": "Seguros — Valor em Risco",
    "trib_itbi": "Tributário — Base de Cálculo ITBI",
    "trib_itcmd": "Tributário — Base de Cálculo ITCMD (Herança / Doação)",
    "trib_ir": "Tributário — Imposto de Renda (Ganho de Capital)",
    "trib_iptu_itr": "Tributário — IPTU / ITR Progressivo",
    "inc_registro": "Incorporação — Registro (Lei 4.591/64)",
    "inc_afetacao": "Incorporação — Patrimônio de Afetação (Lei 13.786/2018)",
    "inc_permuta": "Incorporação — Permuta",
    "exec_fid_1": "Execução de Garantia — Alienação Fiduciária 1º Leilão",
    "exec_fid_2": "Execução de Garantia — Alienação Fiduciária 2º Leilão",
    "exec_hipoteca": "Execução de Garantia — Execução Hipotecária",
    "exec_consolidacao": "Execução de Garantia — Consolidação da Propriedade",
    "desap_utilidade": "Desapropriação por Utilidade Pública (Dec.-Lei 3.365/41)",
    "desap_interesse_social": "Desapropriação por Interesse Social (Lei 4.132/62)",
    "desap_reforma_agraria": "Desapropriação — Reforma Agrária (LC 76/93)",
    "reurb_s_e": "Regularização Fundiária — REURB-S / REURB-E (Lei 13.465/2017)",
    "reurb_demarcacao": "Regularização Fundiária — Demarcação Urbanística",
    "outros_contab": "Contabilidade / Balanço Patrimonial (CPC 28 / IFRS 13)",
    "outros_fii": "Fundo de Investimento Imobiliário (FII)",
    "outros_ma": "Fusão e Aquisição (M&A)",
    "outros_bts": "Locação Built to Suit",
    "outros_diligencia": "Due Diligence Imobiliária",
    "compra_venda": "Compra e Venda",
    "financiamento": "Financiamento / Garantia Bancária",
    "judicial": "Uso Judicial / Inventário",
    "inventario": "Inventário / Partilha",
    "locacao": "Locação",
    "garantia": "Garantia de Crédito",
    "permuta": "Permuta",
    "outros": "Outros",
}

RURAL_TYPES = {"rural", "fazenda", "sitio", "chacara", "terreno_rural"}


def _is_rural(ptam: dict) -> bool:
    return str(ptam.get("property_type", "")).lower() in RURAL_TYPES


# ── renderers de seção ────────────────────────────────────────────────────────

def _render_cover(doc: Document, ptam: dict, user: dict) -> None:
    """Capa com logo, título, número, endereço, solicitante, data."""
    company_logo = user.get("_company_logo_bytes")
    if company_logo:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(io.BytesIO(company_logo), width=Inches(1.5))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PARECER TÉCNICO DE AVALIAÇÃO MERCADOLÓGICA")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Avaliação de Imóvel — ABNT NBR 14653")
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = DARK

    doc.add_paragraph()

    numero = ptam.get("numero_ptam") or ptam.get("number") or "—"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PTAM nº {numero}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK

    doc.add_paragraph()

    endereco = ptam.get("property_address") or ptam.get("property_label") or ""
    cidade = ptam.get("property_city") or ""
    estado = ptam.get("property_state") or ""
    cidade_uf = f"{cidade} — {estado}" if cidade and estado else (cidade or estado)
    if endereco or cidade_uf:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = [l for l in [endereco, cidade_uf] if l]
        run = p.add_run("\n".join(lines))
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = DARK

    doc.add_paragraph()

    # Solicitante (novo campo prioritário)
    sol = ptam.get("solicitante_nome") or ptam.get("solicitante") or "Não informado"
    _add_styled_paragraph(doc, f"Solicitante: {sol}", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if cidade_uf:
        _add_styled_paragraph(doc, f"Cidade: {cidade_uf}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    data_ref = ptam.get("conclusion_date") or ptam.get("vistoria_date") or datetime.now().strftime("%d/%m/%Y")
    _add_styled_paragraph(doc, f"Data: {data_ref}", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    company_name = user.get("company", "") or "RomaTec Consultoria Total"
    _add_styled_paragraph(doc, company_name, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          color=GREEN, space_after=Pt(4))

    doc.add_page_break()


def _render_sumario(doc: Document) -> None:
    """Sumário do PTAM — seções + anexos."""
    _add_section_heading(doc, "SUMÁRIO")
    toc_items = [
        ("1",    "Identificação e Objetivo"),
        ("2",    "Documentação Analisada"),
        ("3",    "Identificação do Imóvel"),
        ("4",    "Caracterização do Imóvel"),
        ("5",    "Análise da Região"),
        ("6",    "Homogeneização e Amostras de Mercado"),
        ("7",    "Metodologia"),
        ("8",    "Cálculo de Ponderância"),
        ("9",    "Método de Avaliação / Depreciação"),
        ("10",   "Valor de Avaliação e Resultado"),
        ("11",   "Ressalvas, Pressupostos e Limitações"),
        ("12",   "Base Legal e Normativa"),
        ("13",   "Declaração de Responsabilidade Técnica"),
        ("14",   "Registro Fotográfico"),
        ("15",   "Documentos Anexos"),
        ("A.IV", "Anexo IV — Certidões das Partes (CND)"),
        ("A.V",  "Anexo V — Currículo do Avaliador"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Seção", "Título"]):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1B4D1B")
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
    for num, title in toc_items:
        row = table.add_row().cells
        row[0].text = num
        row[1].text = title
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_page_break()


def _render_identificacao(doc: Document, ptam: dict) -> None:
    """Seção 1 — Identificação e Objetivo."""
    _add_section_heading(doc, "1. IDENTIFICAÇÃO E OBJETIVO")

    numero = ptam.get("numero_ptam") or ptam.get("number") or "—"
    _add_label_value(doc, "Número do PTAM", numero, skip_zero=False)

    # Finalidade
    finalidade_val = ptam.get("finalidade", "")
    finalidade_str = _FINALIDADE_MAP.get(finalidade_val, finalidade_val) or ptam.get("purpose", "")
    if finalidade_str:
        _add_label_value(doc, "Finalidade", finalidade_str)
    if ptam.get("finalidade_outros"):
        _add_label_value(doc, "Especificação", ptam["finalidade_outros"])

    # Solicitante completo
    _add_subsection_heading(doc, "Dados do Solicitante")
    for label, key in [
        ("Nome / Razão Social", "solicitante_nome"),
        ("CPF / CNPJ", "solicitante_cpf_cnpj"),
        ("Endereço", "solicitante_endereco"),
        ("Telefone", "solicitante_telefone"),
        ("E-mail", "solicitante_email"),
    ]:
        _add_label_value(doc, label, ptam.get(key))
    # fallback legacy
    if not ptam.get("solicitante_nome") and ptam.get("solicitante"):
        _add_label_value(doc, "Solicitante", ptam["solicitante"])

    # Dados judiciais (exibe apenas se preenchidos)
    judicial_keys = ["judicial_process", "judicial_action", "forum", "requerente", "requerido", "judge"]
    if any(ptam.get(k) for k in judicial_keys):
        _add_subsection_heading(doc, "Dados do Processo Judicial")
        for label, key in [
            ("Processo", "judicial_process"),
            ("Ação", "judicial_action"),
            ("Fórum", "forum"),
            ("Requerente", "requerente"),
            ("Requerido", "requerido"),
            ("Juiz", "judge"),
        ]:
            _add_label_value(doc, label, ptam.get(key))


def _render_documentos_analisados(doc: Document, ptam: dict) -> None:
    """Seção 2 — Documentação Analisada."""
    docs = ptam.get("documentos_analisados") or []
    fotos_count = len(ptam.get("fotos_imovel") or [])
    doc_count = len(ptam.get("fotos_documentos") or [])

    _add_section_heading(doc, "2. DOCUMENTAÇÃO ANALISADA")

    docs_label = {
        "matricula": "Matrícula do imóvel",
        "IPTU": "Carnê de IPTU",
        "planta": "Planta / Projeto aprovado",
        "escritura": "Escritura / Contrato",
        "fotos": "Fotografias do imóvel",
        "habite_se": "Habite-se / Auto de conclusão",
        "geo_rural": "Georreferenciamento (rural)",
        "outros_docs": "Outros documentos",
    }
    if docs:
        items_text = " | ".join(docs_label.get(d, d) for d in docs)
        _add_label_value(doc, "Documentos", items_text)
    if fotos_count > 0:
        _add_label_value(doc, "Fotos do imóvel", f"{fotos_count} foto(s) anexada(s)")
    if doc_count > 0:
        _add_label_value(doc, "Documentos digitalizados", f"{doc_count} arquivo(s) anexado(s)")
    if not docs and fotos_count == 0 and doc_count == 0:
        _add_styled_paragraph(doc, "Documentação não especificada.")


def _render_imovel(doc: Document, ptam: dict) -> None:
    """Seção 3 — Identificação do Imóvel."""
    _add_section_heading(doc, "3. IDENTIFICAÇÃO DO IMÓVEL")
    rural = _is_rural(ptam)

    if ptam.get("property_label"):
        _add_label_value(doc, "Rótulo / Título", ptam["property_label"])
    for label, key in [
        ("Tipo", "property_type"),
        ("Endereço", "property_address"),
        ("Bairro", "property_neighborhood"),
        ("CEP", "property_cep"),
        ("Matrícula", "property_matricula"),
        ("Cartório / Ofício", "property_cartorio"),
        ("Proprietário", "property_owner"),
        ("Confrontações", "property_confrontations"),
    ]:
        _add_label_value(doc, label, ptam.get(key))

    # Cidade + Estado
    cidade = (ptam.get("property_city") or "").strip()
    estado = (ptam.get("property_state") or "").strip()
    if cidade and estado:
        _add_label_value(doc, "Cidade / UF", f"{cidade} — {estado}")
    elif cidade:
        _add_label_value(doc, "Cidade", cidade)
    elif estado:
        _add_label_value(doc, "Estado", estado)

    # Coordenadas GPS
    if ptam.get("property_gps_lat") and ptam.get("property_gps_lng"):
        _add_label_value(doc, "Coordenadas GPS",
                         f"{ptam['property_gps_lat']}, {ptam['property_gps_lng']}")

    # Áreas
    if rural:
        if ptam.get("property_area_ha"):
            _add_label_value(doc, "Área total", _format_area_ha(ptam["property_area_ha"]))
        if ptam.get("property_area_sqm"):
            _add_label_value(doc, "Área construída / benfeitorias", _format_area(ptam["property_area_sqm"]))
        if ptam.get("perimetro_m"):
            _add_label_value(doc, "Perímetro",
                             f"{float(ptam['perimetro_m']):,.2f} m".replace(",", "X").replace(".", ",").replace("X", "."))
    else:
        if ptam.get("property_area_sqm"):
            _add_label_value(doc, "Área total (m²)", _format_area(ptam["property_area_sqm"]))
        if ptam.get("property_area_ha"):
            _add_label_value(doc, "Área (ha)", _format_area_ha(ptam["property_area_ha"]))

    if ptam.get("property_description"):
        _add_styled_paragraph(doc, ptam["property_description"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Proprietários (tabela dinâmica)
    proprietarios = [p for p in (ptam.get("proprietarios") or [])
                     if isinstance(p, dict) and p.get("nome")]
    if proprietarios:
        _add_subsection_heading(doc, "Proprietário(s) do Imóvel")
        headers = ["Nome / Razão Social", "CPF / CNPJ", "Fração / Percentual"]
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_shading(hdr[i], "1B4D1B")
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr[i].paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
        for p in proprietarios:
            row = table.add_row().cells
            row[0].text = p.get("nome", "")
            row[1].text = p.get("cpf_cnpj", "")
            row[2].text = p.get("percentual", "")
        doc.add_paragraph()

    # Registros rurais
    if rural:
        rural_fields = [
            ("SIGEF — Sistema de Gestão Fundiária", "certificacao_sigef"),
            ("INCRA — Cadastro no INCRA", "cadastro_incra"),
            ("CCIR — Certificado de Cadastro de Imóvel Rural", "ccir"),
            ("NIRF / CIB", "nirf_cib"),
            ("CAR — Cadastro Ambiental Rural", "car"),
        ]
        rural_items = [(l, ptam.get(k)) for l, k in rural_fields if ptam.get(k)]
        if rural_items:
            _add_subsection_heading(doc, "Registros Rurais")
            for l, v in rural_items:
                _add_label_value(doc, l, v)


def _render_caracterizacao(doc: Document, ptam: dict) -> None:
    """Seção 4 — Caracterização do Imóvel (campos físicos/construtivos)."""
    keys = ["imovel_area_terreno", "imovel_area_construida", "imovel_area_a_considerar",
            "imovel_idade", "imovel_estado_conservacao", "imovel_padrao_acabamento",
            "imovel_num_quartos", "imovel_num_banheiros", "imovel_num_vagas",
            "imovel_piscina", "imovel_caracteristicas_adicionais"]
    if not any(ptam.get(k) for k in keys):
        return

    _add_section_heading(doc, "4. CARACTERIZAÇÃO DO IMÓVEL")

    if ptam.get("imovel_area_terreno"):
        _add_label_value(doc, "Área do Terreno", _format_area(ptam["imovel_area_terreno"]))
    if ptam.get("imovel_area_construida"):
        _add_label_value(doc, "Área Construída", _format_area(ptam["imovel_area_construida"]))
    if ptam.get("imovel_area_a_considerar") is not None:
        _add_label_value(doc, "Área Considerada no Cálculo (*)", _format_area(ptam["imovel_area_a_considerar"]))
        _add_styled_paragraph(
            doc,
            "* Área efetiva utilizada para o cálculo do valor total de avaliação.",
            italic=True,
            color=RGBColor(85, 85, 85),
        )
    if ptam.get("imovel_idade"):
        _add_label_value(doc, "Idade Aproximada", f"{ptam['imovel_idade']} anos")

    estado_map = {"otimo": "Ótimo", "bom": "Bom", "regular": "Regular",
                  "ruim": "Ruim", "pessimo": "Péssimo"}
    if ptam.get("imovel_estado_conservacao"):
        estado_label = estado_map.get(ptam["imovel_estado_conservacao"], ptam["imovel_estado_conservacao"])
        _add_label_value(doc, "Estado de Conservação", estado_label)

    padrao_map = {"alto": "Alto", "medio": "Médio", "simples": "Simples", "minimo": "Mínimo"}
    if ptam.get("imovel_padrao_acabamento"):
        padrao_label = padrao_map.get(ptam["imovel_padrao_acabamento"], ptam["imovel_padrao_acabamento"])
        _add_label_value(doc, "Padrão de Acabamento", padrao_label)

    comodos = []
    if ptam.get("imovel_num_quartos"):
        comodos.append(f"{ptam['imovel_num_quartos']} quarto(s)")
    if ptam.get("imovel_num_banheiros"):
        comodos.append(f"{ptam['imovel_num_banheiros']} banheiro(s)")
    if ptam.get("imovel_num_vagas"):
        comodos.append(f"{ptam['imovel_num_vagas']} vaga(s) de garagem")
    if comodos:
        _add_label_value(doc, "Cômodos / Dependências", " | ".join(comodos))

    if ptam.get("imovel_piscina"):
        _add_label_value(doc, "Piscina", "Sim")

    if ptam.get("imovel_caracteristicas_adicionais"):
        _add_label_value(doc, "Características Adicionais", ptam["imovel_caracteristicas_adicionais"])


def _render_vistoria(doc: Document, ptam: dict) -> None:
    """Seção 5 — Vistoria Técnica (15 sub-seções PTAM nº 7010)."""
    _add_section_heading(doc, "5. VISTORIA TÉCNICA")

    _add_label_value(doc, "Data da Vistoria", ptam.get("vistoria_date"))
    _add_label_value(doc, "Responsável pela Vistoria", ptam.get("vistoria_responsavel"))
    _add_label_value(doc, "Condições de Acesso", ptam.get("vistoria_condicoes"))

    for label, key in [
        ("Objetivo da Vistoria", "vistoria_objective"),
        ("Metodologia Adotada", "vistoria_methodology"),
        ("Topografia", "topography"),
        ("Solo e Cobertura Vegetal", "soil_vegetation"),
        ("3.1 Uso Atual", "uso_atual"),
        ("3.2 Cobertura Vegetal", "cobertura_vegetal"),
        ("3.3 Hidrografia", "hidrografia"),
        ("Benfeitorias Existentes", "benfeitorias"),
        ("5. Infraestrutura Interna", "infraestrutura_interna"),
        ("Acessibilidade e Infraestrutura", "accessibility"),
        ("Contexto Urbano e Mercadológico", "urban_context"),
        ("Estado Geral de Conservação", "conservation_state"),
        ("9. Situação Fundiária", "situacao_fundiaria"),
        ("10. Passivo Ambiental", "passivo_ambiental"),
        ("11. Potencial Exploratório", "potencial_exploratorio"),
        ("12. Aspectos Legais e Restrições", "aspectos_legais"),
        ("13. Restrições de Uso", "restricoes_uso"),
    ]:
        if ptam.get(key):
            _add_subsection_heading(doc, label)
            _add_styled_paragraph(doc, ptam[key], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    if ptam.get("vistoria_synthesis"):
        _add_subsection_heading(doc, "14. Síntese Conclusiva da Vistoria")
        _add_styled_paragraph(doc, ptam["vistoria_synthesis"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _render_regiao(doc: Document, ptam: dict) -> None:
    """Seção 6 — Análise da Região."""
    regiao_keys = ["zoneamento", "regiao_infraestrutura", "regiao_servicos_publicos",
                   "regiao_uso_predominante", "regiao_padrao_construtivo",
                   "regiao_tendencia_mercado", "regiao_observacoes"]
    if not any(ptam.get(k) for k in regiao_keys):
        return

    _add_section_heading(doc, "6. ANÁLISE DA REGIÃO")
    _add_label_value(doc, "Zoneamento (Plano Diretor)", ptam.get("zoneamento"))

    for label, key in [
        ("Infraestrutura Urbana", "regiao_infraestrutura"),
        ("Serviços Públicos", "regiao_servicos_publicos"),
        ("Uso Predominante do Solo", "regiao_uso_predominante"),
        ("Padrão Construtivo da Região", "regiao_padrao_construtivo"),
        ("Tendência de Mercado", "regiao_tendencia_mercado"),
        ("Observações Complementares", "regiao_observacoes"),
    ]:
        if ptam.get(key):
            _add_subsection_heading(doc, label)
            _add_styled_paragraph(doc, ptam[key], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _render_market_samples_table(doc, samples: list) -> None:
    """Tabela completa de amostras de mercado modernas (market_samples)."""
    if not samples:
        return
    headers = ["Nº", "Endereço / Bairro", "Área (m²)", "Valor (R$)", "R$/m²", "Tipo", "Fonte", "Data Coleta"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1B4D1B")
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(9)

    for idx, s in enumerate(samples, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        endereco = f"{s.get('address', '')} / {s.get('neighborhood', '')}".strip(" /")
        row[1].text = endereco or "—"
        area = float(s.get("area") or 0)
        row[2].text = f"{area:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        valor = float(s.get("value") or 0)
        row[3].text = _format_currency(valor).replace("R$ ", "")
        vpm = float(s.get("value_per_sqm") or (valor / area if area > 0 else 0))
        row[4].text = f"{vpm:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        tipo_raw = s.get("tipo_amostra") or "oferta"
        row[5].text = "Consolidada" if tipo_raw == "consolidada" else "Oferta"
        row[6].text = s.get("source") or ""
        row[7].text = s.get("collection_date") or ""
        for i, cell in enumerate(row):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i in (1, 6):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_paragraph()


def _render_amostras_legado_table(doc, samples: list) -> None:
    """Tabela de amostras legacy (impact_areas.samples)."""
    if not samples:
        return
    headers = ["Nº", "Bairro / Local", "Área Total (m²)", "Valor (R$)", "R$/m²"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1B4D1B")
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
    for idx, s in enumerate(samples, start=1):
        row = table.add_row().cells
        row[0].text = str(s.get("number") or idx)
        row[1].text = s.get("neighborhood") or ""
        area = float(s.get("area_total") or 0)
        row[2].text = _format_area(area).replace(" m²", "")
        valor = float(s.get("value") or 0)
        row[3].text = _format_currency(valor).replace("R$ ", "")
        vps = s.get("value_per_sqm") or (valor / area if area > 0 else 0)
        row[4].text = f"{float(vps):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _render_mercado(doc: Document, ptam: dict) -> None:
    """Seção 7 — Homogeneização e Amostras de Mercado."""
    market_samples = ptam.get("market_samples") or []
    market_analysis = ptam.get("market_analysis") or ""
    if not market_samples and not market_analysis:
        return

    _add_section_heading(doc, "7. HOMOGENEIZAÇÃO E AMOSTRAS DE MERCADO")

    if market_analysis:
        _add_styled_paragraph(doc, market_analysis, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    if market_samples:
        _add_subsection_heading(doc, "Elementos de Comparação Coletados (NBR 14653-2 item 8.2)")
        _render_market_samples_table(doc, market_samples)

        # Estatísticas
        values = [float(s.get("value_per_sqm") or 0) for s in market_samples if s.get("value_per_sqm")]
        if values:
            n = len(values)
            avg = sum(values) / n
            sorted_v = sorted(values)
            median = (sorted_v[n // 2 - 1] + sorted_v[n // 2]) / 2 if n % 2 == 0 else sorted_v[n // 2]
            variance = sum((v - avg) ** 2 for v in values) / n
            std = variance ** 0.5
            cv = (std / avg * 100) if avg > 0 else 0
            stats = (
                f"Estatísticas: n={n} | Média: {_format_currency(avg)}/m² | "
                f"Mediana: {_format_currency(median)}/m² | "
                f"Desvio Padrão: {_format_currency(std)} | CV: {cv:.2f}%"
            )
            _add_styled_paragraph(doc, stats, alignment=WD_ALIGN_PARAGRAPH.LEFT, size=10)

        # Fotos das amostras
        amostras_com_foto = [s for s in market_samples if s.get("_image_bytes")]
        if amostras_com_foto:
            _add_subsection_heading(doc, "Fotografias das Amostras")
            for idx, s in enumerate(market_samples, start=1):
                img_bytes = s.get("_image_bytes")
                if not img_bytes:
                    continue
                endereco = f"{s.get('address', '')} / {s.get('neighborhood', '')}".strip(" /")
                _add_styled_paragraph(doc, f"Amostra {idx} — {endereco or ''}", bold=True)
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(5))
                    doc.add_paragraph()
                except Exception:
                    _add_styled_paragraph(doc, f"[Foto amostra {idx} — erro ao carregar]")


def _render_metodologia(doc: Document, ptam: dict) -> None:
    """Seção 8 — Metodologia."""
    _add_section_heading(doc, "8. METODOLOGIA")

    if ptam.get("methodology"):
        _add_styled_paragraph(doc, ptam["methodology"], bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if ptam.get("methodology_justification"):
        _add_styled_paragraph(doc, ptam["methodology_justification"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if ptam.get("calc_fatores_homogeneizacao"):
        _add_subsection_heading(doc, "Fatores de Homogeneização Aplicados")
        _add_styled_paragraph(doc, ptam["calc_fatores_homogeneizacao"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if ptam.get("calc_grau_fundamentacao"):
        _add_label_value(doc, "Grau de Fundamentação (NBR 14653-2)", ptam["calc_grau_fundamentacao"])
    if ptam.get("calc_observacoes"):
        _add_label_value(doc, "Observações sobre os Cálculos", ptam["calc_observacoes"])


def _render_ponderancia(doc: Document, ptam: dict) -> None:
    """Seção 9 — Cálculo de Ponderância (se aplicável)."""
    media = ptam.get("ponderancia_media")
    market_samples = ptam.get("market_samples") or []
    if not media and not market_samples:
        return

    _add_section_heading(doc, "9. CÁLCULO DE PONDERÂNCIA")
    _add_styled_paragraph(
        doc,
        "Método comparativo com filtragem de amostras fora da faixa de 50% a 150% da média simples (ABNT NBR 14653-2).",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    lim_inf = ptam.get("ponderancia_limite_inf")
    lim_sup = ptam.get("ponderancia_limite_sup")
    eliminadas = ptam.get("ponderancia_eliminadas") or []
    valor_final = ptam.get("ponderancia_valor_final")

    valid_samples = [s for s in market_samples if float(s.get("area") or 0) > 0 and float(s.get("value") or 0) > 0]
    if valid_samples:
        _add_subsection_heading(doc, "Quadro de Amostras com Classificação")
        headers = ["Nº", "Bairro / Local", "Área (m²)", "Valor (R$)", "R$/m²", "Situação"]
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_shading(hdr[i], "1B4D1B")
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr[i].paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
        for idx, s in enumerate(valid_samples):
            vpm = float(s.get("value_per_sqm") or 0)
            if vpm == 0 and float(s.get("area") or 0) > 0:
                vpm = float(s.get("value") or 0) / float(s.get("area") or 1)
            eliminada = idx in eliminadas
            row = table.add_row().cells
            row[0].text = str(idx + 1)
            row[1].text = f"{s.get('address', '')} / {s.get('neighborhood', '')}".strip(" /") or "—"
            row[2].text = f"{float(s.get('area') or 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row[3].text = _format_currency(s.get("value", 0)).replace("R$ ", "")
            row[4].text = f"{vpm:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row[5].text = "ELIMINADA" if eliminada else "OK"
            if eliminada:
                _set_cell_shading(row[5], "FFEBEE")
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
        doc.add_paragraph()

    if media:
        _add_subsection_heading(doc, "Resultado do Cálculo de Ponderância")
        if lim_inf is not None:
            _add_label_value(doc, "Limite Inferior (50% da média)", _format_currency(lim_inf) + "/m²")
        if lim_sup is not None:
            _add_label_value(doc, "Limite Superior (150% da média)", _format_currency(lim_sup) + "/m²")
        _add_label_value(doc, "Amostras eliminadas", str(len(eliminadas)))
        restantes = len(valid_samples) - len(eliminadas)
        _add_label_value(doc, "Amostras restantes (válidas)", str(restantes))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Média Ponderada Final: {_format_currency(media)}/m²")
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = GREEN

    if valor_final:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Valor do Imóvel Avaliando: {_format_currency(valor_final)}")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = GOLD


_METODO_LABELS = {
    "ross_heidecke": "Método Ross-Heidecke (Depreciação)",
    "linha_reta": "Método da Linha Reta (Depreciação)",
    "fatores_terreno": "Método dos Fatores de Terreno Urbano",
    "nbr_rural": "Método NBR 14653-3 — Rural / VTN / INCRA",
    "renda": "Método da Renda (Capitalização)",
}


def _render_metodo_avaliacao(doc: Document, ptam: dict) -> None:
    """Seção 10 — Método de Avaliação / Depreciação (se aplicável)."""
    metodo = ptam.get("metodo_avaliacao")
    valor_total = ptam.get("valor_total_metodo")
    if not metodo or valor_total is None:
        return

    _add_section_heading(doc, "10. MÉTODO DE AVALIAÇÃO — DEPRECIAÇÃO E VALORIZAÇÃO")
    label = _METODO_LABELS.get(metodo, metodo)
    _add_label_value(doc, "Método utilizado", label)

    params = ptam.get("metodo_params") or {}

    _add_subsection_heading(doc, "Parâmetros Utilizados")
    if metodo == "ross_heidecke":
        _add_label_value(doc, "Valor de Novo", _format_currency(params.get("valor_novo", 0)))
        _add_label_value(doc, "Idade Atual", f"{params.get('idade_atual', 0)} anos")
        _add_label_value(doc, "Vida Útil", f"{params.get('vida_util', 60)} anos")
        _add_label_value(doc, "Estado de Conservação", params.get("estado", "—"))
        _add_label_value(doc, "Coeficiente de Depreciação (Kd)",
                         f"{ptam.get('depreciacao_percentual', 0):.2f}%")
        _add_label_value(doc, "Valor da Depreciação", _format_currency(ptam.get("valor_depreciacao", 0)))
    elif metodo == "linha_reta":
        _add_label_value(doc, "Valor de Novo", _format_currency(params.get("valor_novo", 0)))
        _add_label_value(doc, "Valor Residual (%)", f"{params.get('residual_pct', 20):.1f}%")
        _add_label_value(doc, "Vida Útil", f"{params.get('vida_util', 40)} anos")
        _add_label_value(doc, "Idade Atual", f"{params.get('idade_atual', 0)} anos")
        _add_label_value(doc, "Depreciação Acumulada", _format_currency(ptam.get("valor_depreciacao", 0)))
    elif metodo == "fatores_terreno":
        _add_label_value(doc, "Valor Unitário de Referência (R$/m²)", _format_currency(params.get("valor_unitario", 0)))
        _add_label_value(doc, "Área do Terreno (m²)", f"{params.get('area_terreno', 0):,.2f} m²".replace(",", "."))
        _add_label_value(doc, "Fator Localização", str(params.get("f_loc", "1.00")))
        _add_label_value(doc, "Fator Topografia", str(params.get("f_topo", "1.00")))
        _add_label_value(doc, "Testada (m)", f"{params.get('testada_m', 0)} m")
    elif metodo == "renda":
        _add_label_value(doc, "Renda Mensal", _format_currency(params.get("renda_mensal", 0)))
        _add_label_value(doc, "Taxa de Capitalização", f"{params.get('taxa_cap', 8):.2f}% a.a.")

    _add_subsection_heading(doc, "Resultado Consolidado")
    if ptam.get("valor_terreno_calc"):
        _add_label_value(doc, "Valor do Terreno / Terra Nua", _format_currency(ptam["valor_terreno_calc"]))
    if ptam.get("valor_benfeitoria"):
        _add_label_value(doc, "Benfeitorias (depreciado)", _format_currency(ptam["valor_benfeitoria"]))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run("Valor Total (Terreno + Benfeitorias): ")
    run1.font.bold = True
    run1.font.size = Pt(12)
    run2 = p.add_run(_format_currency(valor_total))
    run2.font.bold = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = GREEN


def _render_areas_impacto(doc: Document, impact_areas: list) -> None:
    """Seção — Áreas de Impacto (legacy desapropriação)."""
    if not impact_areas:
        return
    _add_section_heading(doc, "AVALIAÇÃO DAS ÁREAS DE IMPACTO")
    for idx, area in enumerate(impact_areas, start=1):
        _add_subsection_heading(doc, f"Área {idx}: {area.get('name', f'Área {idx}')} — {area.get('classification', '')}")
        _add_label_value(doc, "Área Impactada", _format_area(area.get("area_sqm", 0)))
        _add_label_value(doc, "Valor Unitário Adotado", _format_currency(area.get("unit_value", 0)) + "/m²")
        total_v = area.get("total_value") or ((area.get("area_sqm") or 0) * (area.get("unit_value") or 0))
        _add_label_value(doc, "Valor Indenizatório", _format_currency(total_v))
        if area.get("majoration_note"):
            _add_label_value(doc, "Majoração Aplicada", area["majoration_note"])
        if area.get("notes"):
            _add_styled_paragraph(doc, area["notes"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if area.get("samples"):
            _add_styled_paragraph(doc, "Amostragem:", bold=True)
            _render_amostras_legado_table(doc, area["samples"])
        doc.add_paragraph()


def _render_consolidacao_table(doc, impact_areas: list) -> float:
    if not impact_areas:
        return 0.0
    headers = ["Área", "Metragem", "Valor Unitário", "Indenização"]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1B4D1B")
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
    total = 0.0
    for a in impact_areas:
        row = table.add_row()
        row.cells[0].text = f"{a.get('name', '')} ({a.get('classification', '')})"
        row.cells[1].text = _format_area(a.get("area_sqm", 0))
        row.cells[2].text = _format_currency(a.get("unit_value", 0)) + "/m²"
        total_value = a.get("total_value") or ((a.get("area_sqm") or 0) * (a.get("unit_value") or 0))
        total += float(total_value or 0)
        row.cells[3].text = _format_currency(total_value)
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = table.add_row()
    tr.cells[0].merge(tr.cells[2])
    tc = tr.cells[0]
    tc.text = ""
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("TOTAL GERAL DA INDENIZAÇÃO")
    r.bold = True
    _set_cell_shading(tc, "E8F5E9")
    tr.cells[3].text = ""
    rp = tr.cells[3].paragraphs[0]
    rv = rp.add_run(_format_currency(total))
    rv.bold = True
    _set_cell_shading(tr.cells[3], "E8F5E9")
    doc.add_paragraph()
    return total


def _render_resultado(doc: Document, ptam: dict, impact_areas: list) -> None:
    """Seção 10 — Valor de Avaliação e Resultado."""
    _add_section_heading(doc, "10. VALOR DE AVALIAÇÃO E RESULTADO")

    if ptam.get("conclusion_text"):
        _add_styled_paragraph(doc, ptam["conclusion_text"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Consolidação de áreas de impacto (legacy)
    total_calc = 0.0
    if impact_areas:
        _add_subsection_heading(doc, "Consolidação das Indenizações")
        total_calc = _render_consolidacao_table(doc, impact_areas)

    # Valor principal
    valor_total    = ptam.get("resultado_valor_total") or ptam.get("total_indemnity") or total_calc or 0
    valor_unitario = ptam.get("resultado_valor_unitario") or 0
    area_calc      = (
        ptam.get("imovel_area_a_considerar")
        or ptam.get("imovel_area_construida")
        or ptam.get("imovel_area_terreno")
        or 0
    )

    # ── Tabela de cálculo: média × área = total ────────────────────────────────
    if valor_unitario and area_calc:
        _add_subsection_heading(doc, "Cálculo do Valor Final")

        headers = ["Componente", "Valor"]
        calc_rows = [
            ["Média Ponderada Final",
             f"{_format_currency(valor_unitario)}/m²"],
            ["Área do Imóvel Avaliando",
             _format_area(area_calc)],
            [f"Valor Final = {_format_currency(valor_unitario)}/m² × {_format_area(area_calc)}",
             _format_currency(valor_total or valor_unitario * area_calc)],
        ]
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            _set_cell_shading(hdr_cells[i], "1B4D1B")
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
        for ridx, (c0, c1) in enumerate(calc_rows):
            row_cells = tbl.add_row().cells
            row_cells[0].text = c0
            row_cells[1].text = c1
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            is_total = ridx == len(calc_rows) - 1
            for cell in row_cells:
                _set_cell_shading(cell, "E8F5E9" if is_total else "FFFFFF")
                for r in cell.paragraphs[0].runs:
                    r.font.bold = is_total
                    r.font.size = Pt(11 if is_total else 10)
        doc.add_paragraph()

    if valor_total:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p.add_run("Valor de Mercado Avaliado: ")
        run1.font.bold = True
        run1.font.size = Pt(12)
        run2 = p.add_run(_format_currency(valor_total))
        run2.font.bold = True
        run2.font.size = Pt(15)
        run2.font.color.rgb = GREEN

        if ptam.get("total_indemnity_words"):
            _add_styled_paragraph(doc, f"({ptam['total_indemnity_words']})",
                                   italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if valor_unitario:
            _add_label_value(doc, "Valor Unitário (R$/m²)", _format_currency(valor_unitario))

    # Intervalo de confiança
    inf_val = ptam.get("resultado_intervalo_inf") or 0
    sup_val = ptam.get("resultado_intervalo_sup") or 0
    if inf_val and sup_val:
        _add_label_value(doc, "Intervalo de Confiança",
                         f"{_format_currency(inf_val)} a {_format_currency(sup_val)}")

    # Grau de Precisão
    grau_map = {"I": "Grau I — Amplitude ≤ 30%", "II": "Grau II — Amplitude ≤ 20%",
                "III": "Grau III — Amplitude ≤ 10%"}
    grau = ptam.get("grau_precisao") or ""
    if grau:
        _add_label_value(doc, "Grau de Precisão (NBR 14653-1 item 9)", grau_map.get(grau, grau))

    # Campo de Arbítrio
    arb_min = ptam.get("campo_arbitrio_min") or 0
    arb_max = ptam.get("campo_arbitrio_max") or 0
    if arb_min or arb_max:
        _add_label_value(doc, "Campo de Arbítrio (NBR 14653-1 item 9.2.4)",
                         f"{_format_currency(arb_min)} a {_format_currency(arb_max)}  (variação máx. ±15%)")

    # Data de referência
    if ptam.get("resultado_data_referencia"):
        _add_label_value(doc, "Data de Referência da Avaliação", ptam["resultado_data_referencia"])

    # Prazo de Validade
    doc.add_paragraph()
    prazo_meses = ptam.get("prazo_validade_meses") or 6
    prazo_str = ptam.get("resultado_prazo_validade") or f"{prazo_meses} meses"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        f"Este Parecer Técnico tem validade de {prazo_str} a contar da data de emissão, "
        "conforme preconiza a ABNT NBR 14653-1. Após esse período, nova avaliação deverá ser realizada."
    ).font.size = Pt(10)


def _render_consideracoes(doc: Document, ptam: dict) -> None:
    """Seção 11 — Ressalvas, Pressupostos e Limitações."""
    has = any(ptam.get(k) for k in ["consideracoes_ressalvas", "consideracoes_pressupostos", "consideracoes_limitacoes"])
    if not has:
        return
    _add_section_heading(doc, "11. RESSALVAS, PRESSUPOSTOS E LIMITAÇÕES")
    for label, key in [
        ("Ressalvas e Limitações", "consideracoes_ressalvas"),
        ("Pressupostos Adotados", "consideracoes_pressupostos"),
        ("Limitações e Advertências", "consideracoes_limitacoes"),
    ]:
        if ptam.get(key):
            _add_subsection_heading(doc, label)
            _add_styled_paragraph(doc, ptam[key], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _render_base_legal(doc: Document) -> None:
    """Seção 12 — Base Legal e Normativa."""
    _add_section_heading(doc, "12. BASE LEGAL E NORMATIVA")
    normas = [
        "ABNT NBR 14653-1: Procedimentos gerais de avaliação de bens.",
        "ABNT NBR 14653-2: Avaliação de imóveis urbanos — Método Comparativo Direto.",
        "ABNT NBR 14653-3: Avaliação de imóveis rurais.",
        "Resolução COFECI 957/2006: Habilita o Corretor de Imóveis para elaboração de PTAM.",
        "Lei 5.194/1966 — Regulamenta as profissões de Engenharia e Arquitetura.",
        "Lei 6.530/1978 — Regulamenta a profissão de Corretor de Imóveis.",
        "Resolução CONFEA 345/90 — Regulamenta a ART.",
        "Lei 13.786/2018 — Alienação e incorporação imobiliária.",
        "CPC art. 156 — Perícia Judicial.",
    ]
    for item in normas:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10)


def _render_responsavel(doc: Document, ptam: dict, user: dict) -> None:
    """Seção 13 — Declaração de Responsabilidade Técnica + Assinatura."""
    _add_section_heading(doc, "13. DECLARAÇÃO DE RESPONSABILIDADE TÉCNICA")

    tipo_prof = ptam.get("tipo_profissional") or user.get("role", "")
    tipo_map = {
        "corretor": "Corretor de Imóveis habilitado nos termos da Resolução COFECI 957/2006",
        "engenheiro": "Engenheiro com Anotação de Responsabilidade Técnica (ART) conforme Resolução CONFEA 345/90",
        "arquiteto": "Arquiteto e Urbanista com Registro de Responsabilidade Técnica (RRT)",
        "perito_judicial": "Perito Judicial cadastrado no Tribunal, habilitado nos termos do CPC art. 156",
        "agronomo": "Engenheiro Agrônomo com Anotação de Responsabilidade Técnica (ART)",
        "tecnico": "Técnico em Transações Imobiliárias (CFT)",
    }
    tipo_desc = tipo_map.get(tipo_prof, tipo_prof or "Profissional habilitado")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        f"O(A) profissional signatário(a) deste Parecer Técnico de Avaliação Mercadológica é "
        f"{tipo_desc}, responsabilizando-se técnica e legalmente pelo conteúdo e pelos valores "
        "aqui expressos, conforme as normas regulamentadoras vigentes."
    ).font.size = Pt(11)

    art_rrt = ptam.get("art_rrt_numero") or ""
    if art_rrt:
        _add_label_value(doc, "Nº ART / RRT (Res. CONFEA 345/90)", art_rrt)

    doc.add_paragraph()

    # Local / data
    city = ptam.get("conclusion_city") or ptam.get("property_city") or ""
    date_str = ptam.get("conclusion_date") or datetime.now().strftime("%d/%m/%Y")
    loc_text = f"{city}, {date_str}." if city else date_str
    loc_p = doc.add_paragraph(loc_text)
    loc_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    sig_p = doc.add_paragraph("_" * 50)
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Nome: prioridade ptam.responsavel_nome > user.name
    nome = ptam.get("responsavel_nome") or user.get("name") or user.get("nome") or ""
    if nome:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(nome)
        run.font.bold = True
        run.font.size = Pt(12)

    role = user.get("role", "")
    if role:
        _add_styled_paragraph(doc, role, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Registros: CRECI/CNAI do PTAM, depois fallback do user
    creci = ptam.get("responsavel_creci") or user.get("creci") or user.get("crea") or ""
    cnai = ptam.get("responsavel_cnai") or user.get("cna") or ""
    registro = ptam.get("registro_profissional") or ""

    if creci:
        _add_styled_paragraph(doc, creci, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if cnai:
        _add_styled_paragraph(doc, cnai, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if registro:
        _add_styled_paragraph(doc, registro, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if art_rrt:
        _add_styled_paragraph(doc, f"ART/RRT nº {art_rrt}", alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _render_fotos(doc: Document, fotos: list) -> None:
    """Seção 14 — Registro Fotográfico (fotos grandes, 2 por página)."""
    if not fotos:
        return

    _add_section_heading(doc, "14. REGISTRO FOTOGRÁFICO")
    _add_styled_paragraph(doc, "Fotografias do imóvel avaliado, obtidas na data da vistoria:")
    _add_styled_paragraph(doc, f"Total de {len(fotos)} foto(s) anexada(s) ao processo.")
    doc.add_paragraph()

    for i, foto in enumerate(fotos[:20]):
        img_bytes = None
        caption = f"Foto {i + 1}"

        if isinstance(foto, dict):
            img_bytes = foto.get("_image_bytes")
            caption = (
                foto.get("description") or foto.get("descricao")
                or foto.get("caption") or foto.get("legenda")
                or f"Foto {i + 1}"
            )
            if not img_bytes:
                import base64, urllib.request
                url = foto.get("url", "")
                if url:
                    try:
                        if url.startswith("data:"):
                            img_bytes = base64.b64decode(url.split(",", 1)[-1])
                        else:
                            with urllib.request.urlopen(url, timeout=5) as resp:
                                img_bytes = resp.read()
                    except Exception:
                        img_bytes = None

        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_caption.add_run(caption)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = GREEN

        if img_bytes:
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
            except Exception:
                _add_styled_paragraph(doc, f"[{caption} — erro ao carregar imagem]",
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _add_styled_paragraph(doc, f"[{caption} — indisponível]",
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

        # Quebra de página a cada 2 fotos
        if (i + 1) % 2 == 0 and i < len(fotos) - 1:
            doc.add_page_break()


def _render_documentos(doc: Document, docs: list) -> None:
    """Seção 15 — Documentos Anexos."""
    if not docs:
        return

    _add_section_heading(doc, "15. DOCUMENTOS ANEXOS")
    _add_styled_paragraph(doc, "Certidões, escrituras e demais documentos digitalizados anexados a este laudo:")
    doc.add_paragraph()

    for i, doc_item in enumerate(docs[:10]):
        if isinstance(doc_item, dict):
            doc_name = (
                doc_item.get("name") or doc_item.get("nome")
                or doc_item.get("filename") or f"Documento {i + 1}"
            )
            doc_bytes = doc_item.get("_doc_bytes")
            content_type = doc_item.get("content_type", "")
        else:
            doc_name = f"Documento {i + 1}"
            doc_bytes = None
            content_type = ""

        is_pdf = "pdf" in content_type.lower() or doc_name.lower().endswith(".pdf")
        is_image = (
            any(ext in content_type.lower() for ext in ["jpeg", "jpg", "png", "image"])
            or any(doc_name.lower().endswith(ext) for ext in [".jpg", ".jpeg", ".png"])
        )
        tipo_indicador = "[PDF]" if is_pdf else ("[IMG]" if is_image else "[DOC]")

        p = doc.add_paragraph()
        run = p.add_run(f"{i + 1}. {doc_name} {tipo_indicador}")
        run.font.bold = True
        run.font.size = Pt(11)

        if doc_bytes and is_image:
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(io.BytesIO(doc_bytes), width=Inches(6.0))
                doc.add_paragraph()
            except Exception as exc:
                _add_styled_paragraph(doc, f"   [Erro ao exibir imagem: {exc}]", size=9)
        elif doc_bytes and is_pdf:
            tamanho_kb = len(doc_bytes) / 1024
            _add_styled_paragraph(doc, f"   Arquivo PDF ({tamanho_kb:.1f} KB) — disponível nos arquivos digitais.",
                                   size=9, italic=True)
        doc.add_paragraph()


def _render_avaliacoes_areas(doc: Document, ptam: dict) -> None:
    """Seções 9/10 — Avaliação Área 01 e Área 02 (PTAM nº 7010)."""
    area_01_tipo = ptam.get("area_01_tipo") or ""
    area_01_dados = ptam.get("area_01_dados") or ""
    area_01_valor = ptam.get("area_01_valor")
    area_02_tipo = ptam.get("area_02_tipo") or ""
    area_02_dados = ptam.get("area_02_dados") or ""
    area_02_valor = ptam.get("area_02_valor")

    has_01 = area_01_tipo or area_01_dados or area_01_valor
    has_02 = area_02_tipo or area_02_dados or area_02_valor

    if not has_01 and not has_02:
        return

    if has_01:
        _add_section_heading(doc, "9. AVALIAÇÃO — ÁREA 01")
        if area_01_tipo:
            _add_label_value(doc, "Tipo de Área", area_01_tipo)
        if area_01_dados:
            _add_styled_paragraph(doc, area_01_dados, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if area_01_valor is not None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Valor da Área 01: {_format_currency(area_01_valor)}")
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = GREEN
        doc.add_paragraph()

    if has_02:
        _add_section_heading(doc, "10. AVALIAÇÃO — ÁREA 02")
        if area_02_tipo:
            _add_label_value(doc, "Tipo de Área", area_02_tipo)
        if area_02_dados:
            _add_styled_paragraph(doc, area_02_dados, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if area_02_valor is not None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Valor da Área 02: {_format_currency(area_02_valor)}")
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = GREEN
        doc.add_paragraph()

    # Consolidated table
    if has_01 and has_02:
        _add_section_heading(doc, "11. CONSOLIDAÇÃO DAS ÁREAS")
        headers = ["Área", "Tipo", "Valor (R$)"]
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_shading(hdr[i], "1B4D1B")
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr[i].paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
        total = 0.0
        for label, tipo, valor in [
            ("Área 01", area_01_tipo, area_01_valor),
            ("Área 02", area_02_tipo, area_02_valor),
        ]:
            v = float(valor or 0)
            total += v
            row = table.add_row().cells
            row[0].text = label
            row[1].text = tipo or "—"
            row[2].text = _format_currency(v)
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # totals row
        tr = table.add_row().cells
        tr[0].text = "TOTAL GERAL"
        tr[1].text = ""
        tr[2].text = _format_currency(total)
        for cell in tr:
            _set_cell_shading(cell, "E8F5E9")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Valor Total das Áreas: {_format_currency(total)}")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = GOLD
        doc.add_paragraph()


def _render_certidoes(doc: Document, cnd_consultas: list) -> None:
    """Seção Certidões das Partes (CND)."""
    if not cnd_consultas:
        return
    _add_section_heading(doc, "ANEXO IV — CERTIDÕES DAS PARTES (CND)")
    _add_styled_paragraph(
        doc,
        "Certidões Negativas de Débito consultadas junto aos órgãos oficiais para as partes envolvidas nesta avaliação.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    doc.add_paragraph()

    _PROVIDER_LABELS = {
        "receita": "Receita Federal",
        "pgfn": "PGFN - Dívida Ativa",
        "tst": "TST - Certidão Trabalhista",
        "trf1": "TRF1 - Justiça Federal",
        "tjma": "TJMA - Justiça Estadual",
        "cnib": "CNIB - Indisponibilidade",
        "rfb_cadastro": "Situação Cadastral CPF/CNPJ",
    }
    _RESULTADO_LABELS = {
        "negativa": "Negativa",
        "positiva": "Positiva",
        "indisponivel": "Indisponível",
        "erro": "Erro",
    }

    for item in cnd_consultas:
        consulta = item.get("consulta", {})
        certidoes = item.get("certidoes", [])

        nome = consulta.get("nome_parte", "—")
        cpf_cnpj = consulta.get("cpf_cnpj", "—")
        tipo = consulta.get("tipo_parte", "")
        data_consulta = ""
        if consulta.get("created_at"):
            try:
                from datetime import datetime as _dt
                d = consulta["created_at"]
                if isinstance(d, str):
                    d = _dt.fromisoformat(d.replace("Z", "+00:00"))
                data_consulta = d.strftime("%d/%m/%Y")
            except Exception:
                pass

        _add_subsection_heading(doc, f"{nome} — {tipo}" if tipo else nome)
        _add_label_value(doc, "CPF/CNPJ", cpf_cnpj)
        if data_consulta:
            _add_label_value(doc, "Data da Consulta", data_consulta)

        if certidoes:
            headers = ["Órgão / Certidão", "Resultado"]
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr[i].text = h
                _set_cell_shading(hdr[i], "1B4D1B")
                hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in hdr[i].paragraphs[0].runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
            for c in certidoes:
                row = table.add_row().cells
                row[0].text = _PROVIDER_LABELS.get(c.get("provider", ""), c.get("provider", "—"))
                resultado_raw = c.get("resultado") or c.get("status", "indisponivel")
                row[1].text = _RESULTADO_LABELS.get(resultado_raw, resultado_raw.capitalize())
                row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        else:
            _add_styled_paragraph(doc, "Nenhuma certidão retornada para esta parte.")
        doc.add_paragraph()


def _render_curriculo(doc: Document, perfil: dict | None, user: dict) -> None:
    """Seção 'Currículo do Avaliador' — dados completos do profissional."""
    if not perfil:
        return

    _add_section_heading(doc, "ANEXO IV — CURRÍCULO DO AVALIADOR")

    # ── Cabeçalho com nome e registros ───────────────────────────────────────
    nome = perfil.get("nome_completo") or user.get("name", "")
    if nome:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(nome)
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = GREEN

    # Registros profissionais (CRECI, CNAI, etc.)
    registros = perfil.get("registros") or []
    if registros:
        reg_texts = []
        for r in registros:
            tipo = r.get("tipo", "")
            num = r.get("numero", "")
            uf = r.get("uf", "")
            if tipo and num:
                reg_texts.append(f"{tipo} {num}" + (f"/{uf}" if uf else ""))
        if reg_texts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(reg_texts))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

    # Bio/Resumo
    bio = perfil.get("bio_resumo", "")
    if bio:
        doc.add_paragraph()
        _add_styled_paragraph(doc, bio, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        doc.add_paragraph()

    # ── Formação Acadêmica ──────────────────────────────────────────────────
    formacoes = perfil.get("formacoes") or []
    if formacoes:
        _add_subsection_heading(doc, "Formação Acadêmica")
        for f in formacoes:
            tipo = f.get("tipo", "")
            curso = f.get("curso", "")
            inst = f.get("instituicao", "")
            ano = f.get("ano_conclusao", "")
            carga = f.get("carga_horaria", "")
            parts = [p for p in [tipo, curso] if p]
            linha1 = " ".join(parts)
            linha2_parts = [p for p in [inst, f"Ano: {ano}" if ano else "", f"{carga}h" if carga else ""] if p]
            linha2 = " — ".join(linha2_parts)
            if linha1:
                p = doc.add_paragraph()
                p.add_run("• ").font.size = Pt(11)
                r = p.add_run(linha1)
                r.font.bold = True
                r.font.size = Pt(11)
            if linha2:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(0.5)
                r2 = p2.add_run(linha2)
                r2.font.italic = True
                r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(85, 85, 85)
        doc.add_paragraph()

    # ── Experiência Profissional ─────────────────────────────────────────────
    experiencias = perfil.get("experiencias") or []
    if experiencias:
        _add_subsection_heading(doc, "Experiência Profissional")
        for exp in experiencias:
            cargo = exp.get("cargo", "")
            empresa = exp.get("empresa", "")
            inicio = exp.get("periodo_inicio", "")
            fim = exp.get("periodo_fim", "")
            desc = exp.get("descricao", "")
            periodo = f"{inicio}" + (f" até {fim}" if fim else " — atual")
            if cargo and empresa:
                p = doc.add_paragraph()
                p.add_run("• ").font.size = Pt(11)
                r = p.add_run(f"{cargo} — {empresa}")
                r.font.bold = True
                r.font.size = Pt(11)
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(0.5)
                r2 = p2.add_run(periodo)
                r2.font.italic = True
                r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(85, 85, 85)
            if desc:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Cm(0.5)
                p3.add_run(desc).font.size = Pt(9)
        doc.add_paragraph()

    # ── Especializações e Habilitações ───────────────────────────────────────
    especializacoes = perfil.get("especializacoes") or []
    habilitacoes = perfil.get("habilitacoes") or []
    if especializacoes or habilitacoes:
        _add_subsection_heading(doc, "Especializações e Habilitações")
        if especializacoes:
            p = doc.add_paragraph()
            r1 = p.add_run("Especializações: ")
            r1.font.bold = True
            r1.font.size = Pt(11)
            p.add_run("; ".join(especializacoes)).font.size = Pt(11)
        if habilitacoes:
            p = doc.add_paragraph()
            r1 = p.add_run("Habilitações: ")
            r1.font.bold = True
            r1.font.size = Pt(11)
            p.add_run("; ".join(habilitacoes)).font.size = Pt(11)
        doc.add_paragraph()

    # ── Associações Profissionais ────────────────────────────────────────────
    associacoes = perfil.get("membro_associacoes") or []
    if associacoes:
        _add_subsection_heading(doc, "Associações Profissionais")
        for assoc in associacoes:
            p = doc.add_paragraph()
            p.add_run("• ").font.size = Pt(11)
            p.add_run(assoc).font.size = Pt(11)
        doc.add_paragraph()

    # ── Áreas de Atuação ─────────────────────────────────────────────────────
    areas = perfil.get("areas_atuacao") or []
    if areas:
        _add_subsection_heading(doc, "Áreas de Atuação")
        _add_styled_paragraph(doc, "; ".join(areas))
        doc.add_paragraph()

    # ── Tribunais e Bancos Habilitados ───────────────────────────────────────
    tribunais = perfil.get("tribunais_cadastrado") or []
    bancos = perfil.get("bancos_habilitado") or []
    if tribunais or bancos:
        _add_subsection_heading(doc, "Cadastros e Habilitações")
        if tribunais:
            p = doc.add_paragraph()
            r1 = p.add_run("Tribunais cadastrado: ")
            r1.font.bold = True
            r1.font.size = Pt(11)
            p.add_run(", ".join(tribunais)).font.size = Pt(11)
        if bancos:
            p = doc.add_paragraph()
            r1 = p.add_run("Bancos habilitado: ")
            r1.font.bold = True
            r1.font.size = Pt(11)
            p.add_run(", ".join(bancos)).font.size = Pt(11)
        doc.add_paragraph()

    # ── Contato ──────────────────────────────────────────────────────────────
    _add_subsection_heading(doc, "Contato")
    contato_items = []
    email_prof = perfil.get("email_profissional") or user.get("email", "")
    telefone = perfil.get("telefone", "")
    site = perfil.get("site", "")
    if email_prof:
        contato_items.append(f"E-mail: {email_prof}")
    if telefone:
        contato_items.append(f"Telefone: {telefone}")
    if site:
        contato_items.append(f"Site: {site}")
    for item in contato_items:
        p = doc.add_paragraph()
        p.add_run("• ").font.size = Pt(11)
        p.add_run(item).font.size = Pt(11)

    # ── Endereço do Escritório ───────────────────────────────────────────────
    endereco = perfil.get("endereco_escritorio", "")
    cidade = perfil.get("cidade", "")
    uf = perfil.get("uf", "")
    cep = perfil.get("cep", "")
    if any([endereco, cidade, uf]):
        doc.add_paragraph()
        _add_subsection_heading(doc, "Endereço Profissional")
        end_parts = [p for p in [endereco, f"{cidade}/{uf}" if cidade or uf else "", f"CEP: {cep}" if cep else ""] if p]
        if end_parts:
            _add_styled_paragraph(doc, " — ".join(end_parts))

    # ── Empresa ──────────────────────────────────────────────────────────────
    empresa_nome = perfil.get("empresa_nome", "")
    empresa_cnpj = perfil.get("empresa_cnpj", "")
    if empresa_nome:
        doc.add_paragraph()
        _add_subsection_heading(doc, "Empresa")
        p = doc.add_paragraph()
        r = p.add_run(empresa_nome)
        r.font.bold = True
        r.font.size = Pt(11)
        if empresa_cnpj:
            p.add_run(f" — CNPJ: {empresa_cnpj}").font.size = Pt(11)

    # ── Estatísticas ─────────────────────────────────────────────────────────
    num_laudos = perfil.get("numero_laudos_emitidos", 0)
    if num_laudos:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Total de laudos emitidos: {num_laudos}")
        r.font.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(102, 102, 102)


# ── entry point ───────────────────────────────────────────────────────────────

def generate_ptam_docx(ptam: dict, user: dict, cnd_consultas: list | None = None, perfil_avaliador: dict | None = None) -> bytes:
    """Gera DOCX completo do PTAM com paridade total ao PDF. Retorna bytes."""
    if user is None:
        user = {}

    doc = Document()

    # Margens A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    impact_areas = ptam.get("impact_areas") or []

    # ── Capa ──────────────────────────────────────────────────────────────────
    _render_cover(doc, ptam, user)

    # ── Sumário ───────────────────────────────────────────────────────────────
    _render_sumario(doc)

    # ── Seção 1: Identificação e Objetivo ────────────────────────────────────
    _render_identificacao(doc, ptam)
    doc.add_paragraph()

    # ── Seção 2: Documentação Analisada ──────────────────────────────────────
    _render_documentos_analisados(doc, ptam)
    doc.add_paragraph()

    # ── Seção 3: Identificação do Imóvel ─────────────────────────────────────
    _render_imovel(doc, ptam)
    doc.add_paragraph()

    # ── Seção 4: Caracterização do Imóvel ─────────────────────────────────────
    _render_caracterizacao(doc, ptam)
    doc.add_paragraph()

    # ── Seção 5: Análise da Região ────────────────────────────────────────────
    _render_regiao(doc, ptam)
    doc.add_paragraph()

    # ── Seção 6: Amostras de Mercado ──────────────────────────────────────────
    _render_mercado(doc, ptam)
    doc.add_paragraph()

    # ── Seção 7: Metodologia ──────────────────────────────────────────────────
    _render_metodologia(doc, ptam)
    doc.add_paragraph()

    # ── Seção 8: Ponderância ──────────────────────────────────────────────────
    _render_ponderancia(doc, ptam)
    doc.add_paragraph()

    # ── Seção 9: Método de Avaliação / Depreciação ────────────────────────────
    _render_metodo_avaliacao(doc, ptam)
    doc.add_paragraph()

    # ── Legacy: Áreas de Impacto ───────────────────────────────────────────────
    _render_areas_impacto(doc, impact_areas)
    if impact_areas:
        doc.add_paragraph()

    # ── Seção 10: Resultado da Avaliação ──────────────────────────────────────
    doc.add_page_break()
    _render_resultado(doc, ptam, impact_areas)
    doc.add_paragraph()

    # ── Seção: Considerações, Ressalvas e Pressupostos ────────────────────────
    _render_consideracoes(doc, ptam)
    doc.add_paragraph()

    # ── Anexo III: Base Legal ─────────────────────────────────────────────────
    _add_section_heading(doc, "ANEXO III — BASE LEGAL E NORMATIVA")
    _render_base_legal(doc)
    doc.add_paragraph()

    # ── Responsável Técnico + Assinatura ──────────────────────────────────────
    doc.add_page_break()
    _render_responsavel(doc, ptam, user)

    # ── Anexo I: Registro Fotográfico ─────────────────────────────────────────
    fotos = ptam.get("fotos_imovel") or ptam.get("photos") or []
    if fotos:
        doc.add_page_break()
        _add_section_heading(doc, "ANEXO I — FICHA DO IMÓVEL, FOTOS E DOCUMENTOS")
        _render_fotos(doc, fotos)

    # ── Documentos Anexos ─────────────────────────────────────────────────────
    docs_list = ptam.get("fotos_documentos") or []
    if docs_list:
        doc.add_page_break()
        _render_documentos(doc, docs_list)

    # ── Anexo IV: Certidões das Partes (CND) ──────────────────────────────────
    if cnd_consultas:
        doc.add_page_break()
        _render_certidoes(doc, cnd_consultas)

    # ── Anexo V: Currículo do Avaliador ───────────────────────────────────────
    if perfil_avaliador:
        doc.add_page_break()
        _render_curriculo(doc, perfil_avaliador, user)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

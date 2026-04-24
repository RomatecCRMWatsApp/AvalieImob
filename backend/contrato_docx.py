# @module contrato_docx — Geração de DOCX para Contratos Imobiliários
"""DOCX generator for Contratos Imobiliarios — layout juridico premium RomaTec.

Layout:
  - Times New Roman 12pt em todo o corpo
  - Margens: 3cm esquerda/superior, 2cm direita/inferior (padrao juridico)
  - Logo no cabecalho (reusar padrao ptam_docx.py)
  - Qualificacao completa das partes
  - Clausulas em romano (I, II, III...)
  - Texto justificado com recuo
  - Bloco de assinaturas
  - Recibo de arras como documento separado

Referencias legais:
  CC/2002 art. 104, 481-533 | Lei 8.245/91 | Lei 9.514/97
  Lei 13.786/2018 | COFECI 957/2006 | CPC art. 221
"""
import io
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Cores RomaTec ─────────────────────────────────────────────────────────────
GREEN = RGBColor(27, 77, 27)        # #1B4D1B
GOLD = RGBColor(212, 168, 48)       # #D4A830
WHITE = RGBColor(255, 255, 255)
DARK = RGBColor(26, 26, 26)
GRAY = RGBColor(85, 85, 85)
LIGHT_GREEN = RGBColor(232, 245, 233)

# Fonte juridica padrao
FONT_JURIDICA = "Times New Roman"


# ── Helpers de celula ─────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "4")
        edge_el.set(qn("w:color"), "1B4D1B")
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


# ── Helpers de paragrafo ─────────────────────────────────────────────────────

def _p(doc: Document, text: str = "", bold: bool = False, italic: bool = False,
       size: int = 12, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
       space_before: float = 0, space_after: float = 6,
       left_indent: float = 0, first_line_indent: float = 0) -> None:
    """Adiciona paragrafo com fonte Times New Roman 12pt por padrao."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent:
        pf.left_indent = Cm(left_indent)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        run.font.name = FONT_JURIDICA
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p


def _p_bold_value(doc: Document, label: str, value: str, size: int = 12) -> None:
    """Linha: **label:** valor"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    r1.font.name = FONT_JURIDICA
    r1.font.size = Pt(size)
    r1.font.bold = True
    if value:
        r2 = p.add_run(str(value))
        r2.font.name = FONT_JURIDICA
        r2.font.size = Pt(size)


def _section_heading(doc: Document, text: str) -> None:
    """Cabecalho de secao: tabela 1x1 verde escuro, texto branco, Times New Roman."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.text = ""
    _set_cell_shading(cell, "1B4D1B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_JURIDICA
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _clausula_heading(doc: Document, numero_romano: str, titulo: str) -> None:
    """Titulo de clausula: CLAUSULA I — TITULO"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"CLÁUSULA {numero_romano} — {titulo.upper()}")
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK


def _format_currency(value) -> str:
    try:
        v = float(value) if value else 0
        return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "R$ 0,00"


def _fmt(val) -> str:
    if val is None:
        return ""
    return str(val).strip()


_ROMANOS = [
    (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
    (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
    (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
]


def _to_roman(n: int) -> str:
    if n <= 0:
        return str(n)
    result = ""
    for val, sym in _ROMANOS:
        while n >= val:
            result += sym
            n -= val
    return result


# ── Secoes do contrato ────────────────────────────────────────────────────────

def _render_cover(doc: Document, contrato: dict, user: dict) -> None:
    """Capa do contrato: logo, titulo, numero, partes principais."""
    # Logo
    logo_bytes = user.get("_company_logo_bytes")
    if logo_bytes:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.5))
        except Exception:
            pass
    else:
        company_name = user.get("company") or "RomaTec Consultoria Total"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company_name)
        run.font.name = FONT_JURIDICA
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = GREEN

    doc.add_paragraph()

    # Titulo principal
    tipo = contrato.get("tipo_contrato") or "CONTRATO IMOBILIÁRIO"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tipo.upper())
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = GREEN

    doc.add_paragraph()

    # Numero do contrato
    numero = contrato.get("numero_contrato") or "—"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Nº {numero}")
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = DARK

    doc.add_paragraph()

    # Partes resumidas
    partes = contrato.get("partes") or []
    for p_item in partes[:4]:
        qual = p_item.get("qualificacao", "PARTE")
        if p_item.get("tipo") == "pf" and p_item.get("pf"):
            nome = p_item["pf"].get("nome", "")
        elif p_item.get("tipo") == "pj" and p_item.get("pj"):
            nome = p_item["pj"].get("razao_social", "")
        else:
            nome = ""
        if nome:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p.add_run(f"{qual.upper()}: ")
            run1.font.name = FONT_JURIDICA
            run1.font.size = Pt(12)
            run1.font.bold = True
            run2 = p.add_run(nome)
            run2.font.name = FONT_JURIDICA
            run2.font.size = Pt(12)

    doc.add_paragraph()

    # Cidade/Data
    cidade = contrato.get("cidade_assinatura") or user.get("city") or ""
    uf = contrato.get("uf_assinatura") or user.get("uf") or ""
    data_ref = contrato.get("data_assinatura") or datetime.now().strftime("%d/%m/%Y")
    if cidade or uf:
        loc = f"{cidade}/{uf}" if cidade and uf else (cidade or uf)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{loc}, {data_ref}")
        run.font.name = FONT_JURIDICA
        run.font.size = Pt(11)

    doc.add_page_break()


def _render_qualificacao_partes(doc: Document, contrato: dict) -> None:
    """Secao: Qualificacao das Partes."""
    _section_heading(doc, "QUALIFICAÇÃO DAS PARTES")

    partes = contrato.get("partes") or []
    if not partes:
        _p(doc, "Nenhuma parte cadastrada.")
        return

    for idx, parte in enumerate(partes):
        qual = parte.get("qualificacao") or f"PARTE {idx + 1}"

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f"{qual.upper()}:")
        run.font.name = FONT_JURIDICA
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = GREEN

        if parte.get("tipo") == "pf" and parte.get("pf"):
            pf = parte["pf"]
            _qualifica_pf(doc, pf)
        elif parte.get("tipo") == "pj" and parte.get("pj"):
            pj = parte["pj"]
            _qualifica_pj(doc, pj)

    doc.add_paragraph()


def _qualifica_pf(doc: Document, pf: dict) -> None:
    """Paragrafo de qualificacao para pessoa fisica — formato juridico padrao."""
    nome = _fmt(pf.get("nome"))
    nac = _fmt(pf.get("nacionalidade")) or "brasileiro(a)"
    prof = _fmt(pf.get("profissao"))
    ec = _fmt(pf.get("estado_civil"))
    cpf = _fmt(pf.get("cpf"))
    rg = _fmt(pf.get("rg"))
    rg_orgao = _fmt(pf.get("rg_orgao"))
    rg_uf = _fmt(pf.get("rg_uf"))
    end = _fmt(pf.get("endereco"))
    cidade = _fmt(pf.get("cidade"))
    uf = _fmt(pf.get("uf"))
    cep = _fmt(pf.get("cep"))

    rg_str = rg
    if rg_orgao:
        rg_str += f" {rg_orgao}"
    if rg_uf:
        rg_str += f"/{rg_uf}"

    partes_qual = []
    if nome:
        partes_qual.append(nome)
    if nac:
        partes_qual.append(nac)
    if ec:
        partes_qual.append(ec)
    if prof:
        partes_qual.append(prof)
    if cpf:
        partes_qual.append(f"CPF nº {cpf}")
    if rg_str:
        partes_qual.append(f"RG nº {rg_str}")

    end_parts = [p for p in [end, cidade, uf] if p]
    if end_parts:
        end_str = ", ".join(end_parts)
        if cep:
            end_str += f", CEP {cep}"
        partes_qual.append(f"residente e domiciliado(a) na {end_str}")

    texto = ", ".join(partes_qual) + ("." if partes_qual else "")

    # Conjuge (se casado)
    if pf.get("conjuge_nome"):
        conj_nome = pf.get("conjuge_nome", "")
        conj_cpf = pf.get("conjuge_cpf", "")
        regime = pf.get("regime_bens", "")
        conj_str = f", casado(a) com {conj_nome}"
        if conj_cpf:
            conj_str += f", CPF {conj_cpf}"
        if regime:
            conj_str += f", sob o regime de {regime}"
        texto = texto.rstrip(".") + conj_str + "."

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(texto)
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(12)

    # Procurador
    if pf.get("procurador_nome"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.left_indent = Cm(1)
        run2 = p2.add_run(
            f"Representado(a) por seu procurador(a) {pf['procurador_nome']}, "
            f"CPF {pf.get('procurador_cpf', '')}, "
            f"conforme {pf.get('procurador_instrumento', 'instrumento de mandato anexo')}."
        )
        run2.font.name = FONT_JURIDICA
        run2.font.size = Pt(12)
        run2.font.italic = True


def _qualifica_pj(doc: Document, pj: dict) -> None:
    """Paragrafo de qualificacao para pessoa juridica — formato juridico padrao."""
    razao = _fmt(pj.get("razao_social"))
    cnpj = _fmt(pj.get("cnpj"))
    end = _fmt(pj.get("endereco"))
    cidade = _fmt(pj.get("cidade"))
    uf = _fmt(pj.get("uf"))
    cep = _fmt(pj.get("cep"))
    repr_nome = _fmt(pj.get("representante_nome"))
    repr_cpf = _fmt(pj.get("representante_cpf"))
    repr_cargo = _fmt(pj.get("representante_cargo"))

    parts_qual = []
    if razao:
        parts_qual.append(razao)
    if cnpj:
        parts_qual.append(f"inscrita no CNPJ sob o nº {cnpj}")
    end_parts = [p for p in [end, cidade, uf] if p]
    if end_parts:
        end_str = ", ".join(end_parts)
        if cep:
            end_str += f", CEP {cep}"
        parts_qual.append(f"com sede na {end_str}")

    texto = ", ".join(parts_qual) + ("." if parts_qual else "")

    if repr_nome:
        texto = texto.rstrip(".")
        texto += f", neste ato representada por {repr_nome}"
        if repr_cpf:
            texto += f", CPF {repr_cpf}"
        if repr_cargo:
            texto += f", {repr_cargo}"
        texto += "."

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(texto)
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(12)


def _render_objeto(doc: Document, contrato: dict) -> None:
    """Secao de identificacao do objeto do contrato."""
    objeto = contrato.get("objeto")
    if not objeto:
        return

    _section_heading(doc, "DO OBJETO DO CONTRATO")

    tipo_obj = objeto.get("tipo", "imovel_urbano")
    end = _fmt(objeto.get("endereco"))
    complemento = _fmt(objeto.get("complemento"))
    bairro = _fmt(objeto.get("bairro"))
    cidade = _fmt(objeto.get("cidade"))
    uf = _fmt(objeto.get("uf"))
    cep = _fmt(objeto.get("cep"))
    matricula = _fmt(objeto.get("matricula"))
    cartorio = _fmt(objeto.get("cartorio"))
    area_t = objeto.get("area_terreno")
    area_c = objeto.get("area_construida")
    situacao = _fmt(objeto.get("situacao"))
    onus = _fmt(objeto.get("onus"))
    ocupacao = _fmt(objeto.get("ocupacao"))

    if tipo_obj in ("imovel_urbano", "imovel_rural"):
        partes_end = [p for p in [end, complemento, bairro, cidade, uf] if p]
        end_completo = ", ".join(partes_end)
        if cep:
            end_completo += f", CEP {cep}"

        texto_obj = "O presente contrato tem como objeto o imóvel"
        if end_completo:
            texto_obj += f" situado na {end_completo}"
        if matricula:
            texto_obj += f", com matrícula nº {matricula}"
        if cartorio:
            texto_obj += f" no {cartorio}"
        if area_t:
            texto_obj += f", com área de terreno de {float(area_t):,.2f} m²".replace(",", "X").replace(".", ",").replace("X", ".")
        if area_c:
            texto_obj += f" e área construída de {float(area_c):,.2f} m²".replace(",", "X").replace(".", ",").replace("X", ".")
        texto_obj += "."

        _p(doc, texto_obj, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent=0)

        if situacao:
            _p(doc, f"O imóvel encontra-se {situacao}.", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if onus:
            _p(doc, f"Ônus reais: {onus}.", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if ocupacao:
            _p(doc, f"Situação de ocupação: {ocupacao}.", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Campos rurais
        ccir = _fmt(objeto.get("ccir"))
        nirf = _fmt(objeto.get("nirf"))
        car = _fmt(objeto.get("car"))
        denom = _fmt(objeto.get("denominacao"))
        if denom:
            _p_bold_value(doc, "Denominação", denom)
        if ccir:
            _p_bold_value(doc, "CCIR", ccir)
        if nirf:
            _p_bold_value(doc, "NIRF/ITR", nirf)
        if car:
            _p_bold_value(doc, "CAR", car)

        iptu = _fmt(objeto.get("inscricao_iptu"))
        if iptu:
            _p_bold_value(doc, "Inscrição IPTU", iptu)

    elif tipo_obj == "veiculo":
        marca = _fmt(objeto.get("veiculo_marca"))
        modelo = _fmt(objeto.get("veiculo_modelo"))
        ano_fab = objeto.get("veiculo_ano_fabricacao")
        ano_mod = objeto.get("veiculo_ano_modelo")
        placa = _fmt(objeto.get("veiculo_placa"))
        renavam = _fmt(objeto.get("veiculo_renavam"))
        chassi = _fmt(objeto.get("veiculo_chassi"))
        cor = _fmt(objeto.get("veiculo_cor"))
        km = objeto.get("veiculo_km")

        texto_veic = "O presente contrato tem como objeto o veículo"
        partes_veic = [p for p in [marca, modelo] if p]
        if partes_veic:
            texto_veic += f" {' '.join(partes_veic)}"
        if ano_fab or ano_mod:
            texto_veic += f", ano fabricação {ano_fab or '—'}/modelo {ano_mod or '—'}"
        if cor:
            texto_veic += f", cor {cor}"
        if placa:
            texto_veic += f", placa {placa}"
        if renavam:
            texto_veic += f", RENAVAM {renavam}"
        if chassi:
            texto_veic += f", chassi {chassi}"
        if km is not None:
            texto_veic += f", com {km:,} km rodados".replace(",", ".")
        texto_veic += "."
        _p(doc, texto_veic, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    desc = _fmt(objeto.get("descricao_adicional"))
    if desc:
        _p(doc, desc, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_paragraph()


def _render_clausulas(doc: Document, contrato: dict) -> None:
    """Renderiza as clausulas do contrato em romano."""
    clausulas = contrato.get("clausulas") or []
    if not clausulas:
        return

    _section_heading(doc, "DAS CLÁUSULAS E CONDIÇÕES")

    for idx, cla in enumerate(clausulas):
        num = cla.get("numero") or (idx + 1)
        titulo = _fmt(cla.get("titulo")) or f"CLÁUSULA {_to_roman(num)}"
        conteudo = _fmt(cla.get("conteudo"))
        base_legal = _fmt(cla.get("base_legal"))

        # Numero romano do campo _numero_romano ou calculado
        nr = _fmt(cla.get("_numero_romano")) or _to_roman(num)

        _clausula_heading(doc, nr, titulo)

        if conteudo:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(1.25)
            run = p.add_run(conteudo)
            run.font.name = FONT_JURIDICA
            run.font.size = Pt(12)

        if base_legal:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.25)
            run = p.add_run(f"(Base legal: {base_legal})")
            run.font.name = FONT_JURIDICA
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = GRAY

    doc.add_paragraph()


def _render_corretor(doc: Document, contrato: dict) -> None:
    """Secao de corretagem (se houver corretor cadastrado)."""
    corretor = contrato.get("corretor")
    if not corretor or not corretor.get("nome"):
        return

    _section_heading(doc, "DA INTERMEDIAÇÃO / CORRETAGEM")

    nome = _fmt(corretor.get("nome"))
    creci = _fmt(corretor.get("creci"))
    cpf_cnpj = _fmt(corretor.get("cpf_cnpj"))
    imob = _fmt(corretor.get("imobiliaria"))
    imob_cnpj = _fmt(corretor.get("imobiliaria_cnpj"))
    comissao_pct = corretor.get("comissao_percentual")
    comissao_val = corretor.get("comissao_valor")
    responsavel = _fmt(corretor.get("comissao_responsavel"))
    exclusividade = corretor.get("exclusividade", False)
    excl_dias = corretor.get("exclusividade_prazo_dias")

    # Qualificacao do corretor
    texto_corr = f"O presente contrato foi intermediado pelo(a) Corretor(a) de Imóveis {nome}"
    if creci:
        texto_corr += f", inscrito(a) no CRECI sob o nº {creci}"
    if cpf_cnpj:
        texto_corr += f", CPF/CNPJ nº {cpf_cnpj}"
    if imob:
        texto_corr += f", vinculado(a) à imobiliária {imob}"
        if imob_cnpj:
            texto_corr += f", CNPJ {imob_cnpj}"
    texto_corr += ", nos termos dos art. 722 a 729 do Código Civil e da Resolução COFECI 957/2006."

    _p(doc, texto_corr, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Comissao
    if comissao_pct or comissao_val:
        txt_com = "A comissão de corretagem"
        if comissao_pct:
            txt_com += f" é de {comissao_pct}%"
        if comissao_val:
            txt_com += f" ({_format_currency(comissao_val)})"
        if responsavel:
            txt_com += f", a ser paga pelo(a) {responsavel}"
        txt_com += ", conforme art. 724 CC."
        _p(doc, txt_com, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Exclusividade
    if exclusividade and excl_dias:
        _p(doc,
           f"O presente contrato de corretagem é EXCLUSIVO pelo prazo de {excl_dias} dias, "
           "sendo devida a comissão mesmo que o negócio seja concluído diretamente pelas partes "
           "durante a vigência da exclusividade (art. 726 CC).",
           size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_paragraph()


def _render_pagamento(doc: Document, contrato: dict) -> None:
    """Secao de condicoes de pagamento."""
    cond = contrato.get("condicoes_pagamento")
    if not cond:
        return

    _section_heading(doc, "DO PREÇO E DAS CONDIÇÕES DE PAGAMENTO")

    valor = cond.get("valor_total", 0)
    extenso = _fmt(cond.get("valor_total_extenso"))
    forma = _fmt(cond.get("forma_principal"))
    sinal = cond.get("sinal_valor")
    sinal_data = _fmt(cond.get("sinal_data"))
    arras_tipo = _fmt(cond.get("sinal_arras_tipo"))
    multa = cond.get("multa_inadimplemento")
    juros = cond.get("juros_mora")
    correcao = _fmt(cond.get("correcao_monetaria"))
    financ_banco = _fmt(cond.get("financiamento_banco"))
    financ_valor = cond.get("financiamento_valor")
    financ_prazo = cond.get("financiamento_prazo_meses")

    txt_preco = f"O valor total do presente contrato é de {_format_currency(valor)}"
    if extenso:
        txt_preco += f" ({extenso})"
    if forma:
        txt_preco += f", a ser pago mediante {forma}"
    txt_preco += "."
    _p(doc, txt_preco, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Sinal/Arras
    if sinal and float(sinal) > 0:
        arras_mapa = {
            "confirmatórias": "confirmatórias (art. 417 CC — o sinal se integra ao valor do contrato)",
            "penitenciais": "penitenciais (art. 420 CC — a parte que desistir perderá o sinal; "
                           "se o desistente for o recebedor, restituirá em dobro)",
        }
        arras_desc = arras_mapa.get(arras_tipo, arras_tipo or "a título de sinal")
        txt_arras = (
            f"Como sinal e princípio de pagamento, será pago o valor de {_format_currency(sinal)}, "
            f"a título de arras {arras_desc}"
        )
        if sinal_data:
            txt_arras += f", na data de {sinal_data}"
        txt_arras += "."
        _p(doc, txt_arras, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Financiamento
    if financ_banco:
        txt_financ = f"O saldo restante será financiado junto ao {financ_banco}"
        if financ_valor:
            txt_financ += f", no valor de {_format_currency(financ_valor)}"
        if financ_prazo:
            txt_financ += f", em {financ_prazo} meses"
        txt_financ += "."
        _p(doc, txt_financ, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Parcelas
    parcelas = cond.get("parcelas") or []
    if parcelas:
        _p(doc, "QUADRO DE PARCELAS:", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
        headers = ["Nº", "Valor (R$)", "Vencimento", "Forma de Pagamento", "Banco"]
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_shading(hdr[i], "1B4D1B")
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr[i].paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.name = FONT_JURIDICA
                run.font.size = Pt(10)
        for parc in parcelas:
            row = table.add_row().cells
            row[0].text = str(parc.get("numero") or "—")
            row[1].text = _format_currency(parc.get("valor", 0))
            row[2].text = _fmt(parc.get("vencimento"))
            row[3].text = _fmt(parc.get("forma_pagamento"))
            row[4].text = _fmt(parc.get("banco"))
            for cell in row:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.name = FONT_JURIDICA
                    run.font.size = Pt(10)
        doc.add_paragraph()

    # Penalidades
    if multa is not None:
        _p(doc,
           f"Em caso de inadimplemento, incidirá multa contratual de {multa}% sobre o valor do contrato, "
           "acrescida de juros moratórios"
           + (f" de {juros}% ao mês" if juros else "")
           + (f", com correção monetária pelo índice {correcao}" if correcao else "")
           + " (art. 406 CC; art. 52 §1º CDC).",
           size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    obs = _fmt(cond.get("observacoes"))
    if obs:
        _p(doc, obs, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_paragraph()


def _render_assinaturas(doc: Document, contrato: dict, user: dict) -> None:
    """Bloco de assinaturas: local/data, linhas para partes e testemunhas."""
    _section_heading(doc, "DAS ASSINATURAS")

    cidade = contrato.get("cidade_assinatura") or user.get("city") or "___________"
    uf = contrato.get("uf_assinatura") or user.get("uf") or "___"
    data_ass = contrato.get("data_assinatura") or "_____ de _____________ de _______"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{cidade}/{uf}, {data_ass}.")
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Assinaturas das partes
    partes = contrato.get("partes") or []
    for idx, parte in enumerate(partes):
        qual = parte.get("qualificacao") or f"Parte {idx + 1}"
        if parte.get("tipo") == "pf" and parte.get("pf"):
            nome = parte["pf"].get("nome", "")
            cpf = parte["pf"].get("cpf", "")
            doc_str = f"CPF: {cpf}" if cpf else ""
        elif parte.get("tipo") == "pj" and parte.get("pj"):
            nome = parte["pj"].get("razao_social", "")
            cnpj = parte["pj"].get("cnpj", "")
            doc_str = f"CNPJ: {cnpj}" if cnpj else ""
        else:
            nome = ""
            doc_str = ""

        p = doc.add_paragraph("_" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        if nome:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(nome)
            r.font.name = FONT_JURIDICA
            r.font.size = Pt(12)
            r.font.bold = True
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(f"{qual.upper()}" + (f"  —  {doc_str}" if doc_str else ""))
        r3.font.name = FONT_JURIDICA
        r3.font.size = Pt(11)
        r3.font.italic = True
        doc.add_paragraph()

    # Corretor (se presente)
    corretor = contrato.get("corretor")
    if corretor and corretor.get("nome"):
        p = doc.add_paragraph("_" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(corretor.get("nome", ""))
        r.font.name = FONT_JURIDICA
        r.font.size = Pt(12)
        r.font.bold = True
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        creci = corretor.get("creci", "")
        r3 = p3.add_run("CORRETOR" + (f"  —  CRECI {creci}" if creci else ""))
        r3.font.name = FONT_JURIDICA
        r3.font.size = Pt(11)
        r3.font.italic = True
        doc.add_paragraph()

    # Testemunhas
    testemunhas = contrato.get("testemunhas") or []
    if testemunhas:
        doc.add_paragraph()
        p_test = doc.add_paragraph()
        run = p_test.add_run("TESTEMUNHAS:")
        run.font.name = FONT_JURIDICA
        run.font.size = Pt(12)
        run.font.bold = True

        for idx, test in enumerate(testemunhas[:2]):
            nome_t = _fmt(test.get("nome")) or f"Testemunha {idx + 1}"
            cpf_t = _fmt(test.get("cpf"))

            p = doc.add_paragraph("_" * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(20)
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p2.add_run(nome_t)
            r.font.name = FONT_JURIDICA
            r.font.size = Pt(12)
            r.font.bold = True
            if cpf_t:
                p3 = doc.add_paragraph()
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r3 = p3.add_run(f"CPF: {cpf_t}")
                r3.font.name = FONT_JURIDICA
                r3.font.size = Pt(11)
                r3.font.italic = True
            doc.add_paragraph()


def _render_foro(doc: Document, contrato: dict) -> None:
    """Clausula de foro competente."""
    cidade = contrato.get("cidade_assinatura") or ""
    uf = contrato.get("uf_assinatura") or ""
    loc = f"{cidade}/{uf}" if cidade and uf else (cidade or uf or "desta cidade")

    _section_heading(doc, "DO FORO")
    _p(doc,
       f"Para dirimir quaisquer controversias oriundas do presente contrato, as partes elegem "
       f"o Foro da Comarca de {loc}, renunciando a qualquer outro, por mais privilegiado que seja, "
       "nos termos do art. 63 do Código de Processo Civil.",
       size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_paragraph()


# ── Recibo de Arras ──────────────────────────────────────────────────────────

def generate_recibo_arras_docx(contrato: dict, user: dict) -> bytes:
    """Gera DOCX do Recibo de Sinal/Arras. Retorna bytes."""
    if user is None:
        user = {}

    cond = contrato.get("condicoes_pagamento") or {}
    sinal = float(cond.get("sinal_valor") or 0)
    sinal_extenso = _fmt(cond.get("valor_total_extenso"))
    arras_tipo = _fmt(cond.get("sinal_arras_tipo"))
    sinal_data = _fmt(cond.get("sinal_data")) or datetime.now().strftime("%d/%m/%Y")

    doc = Document()

    # Margens juridicas
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Logo
    logo_bytes = user.get("_company_logo_bytes")
    if logo_bytes:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.3))
        except Exception:
            pass

    doc.add_paragraph()

    # Titulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RECIBO DE SINAL E PRINCÍPIO DE PAGAMENTO")
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = GREEN

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "Arras " + (arras_tipo.capitalize() if arras_tipo else "Confirmatórias")
        + " — art. 417-420 CC"
    )
    run2.font.name = FONT_JURIDICA
    run2.font.size = Pt(11)
    run2.font.italic = True

    doc.add_paragraph()

    # Numero do contrato
    numero = contrato.get("numero_contrato") or "—"
    _p_bold_value(doc, "Contrato nº", numero)

    tipo = contrato.get("tipo_contrato") or ""
    if tipo:
        _p_bold_value(doc, "Tipo de Contrato", tipo)

    doc.add_paragraph()

    # Partes
    _section_heading(doc, "DAS PARTES")
    partes = contrato.get("partes") or []
    for p_item in partes[:2]:
        qual = p_item.get("qualificacao", "Parte")
        if p_item.get("tipo") == "pf" and p_item.get("pf"):
            pf = p_item["pf"]
            nome = pf.get("nome", "")
            cpf = pf.get("cpf", "")
            _p_bold_value(doc, qual, f"{nome}" + (f" — CPF {cpf}" if cpf else ""))
        elif p_item.get("tipo") == "pj" and p_item.get("pj"):
            pj = p_item["pj"]
            razao = pj.get("razao_social", "")
            cnpj = pj.get("cnpj", "")
            _p_bold_value(doc, qual, f"{razao}" + (f" — CNPJ {cnpj}" if cnpj else ""))

    doc.add_paragraph()

    # Valor do sinal
    _section_heading(doc, "DO VALOR DO SINAL")

    txt_valor = (
        f"Declaro que recebi, nesta data, a quantia de {_format_currency(sinal)}"
    )
    if sinal_extenso:
        txt_valor += f" ({sinal_extenso})"
    txt_valor += (
        f", a título de sinal e arras {arras_tipo or 'confirmatórias'}, "
        "correspondente ao princípio de pagamento do contrato acima identificado."
    )
    _p(doc, txt_valor, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Arras legais
    if arras_tipo == "confirmatórias" or not arras_tipo:
        _p(doc,
           "As presentes arras têm natureza CONFIRMATÓRIA, nos termos do art. 417 do Código Civil, "
           "e serão computadas no valor final do contrato, "
           "sem prejudicar o direito de indenização por perdas e danos.",
           size=11, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    elif arras_tipo == "penitenciais":
        _p(doc,
           "As presentes arras têm natureza PENITENCIAL, nos termos do art. 420 do Código Civil. "
           "Em caso de inadimplemento do promitente-comprador, perderá o sinal em benefício do vendedor. "
           "Se o inadimplemento for do promitente-vendedor, este as restituirá em dobro.",
           size=11, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.add_paragraph()

    # Objeto
    objeto = contrato.get("objeto")
    if objeto:
        _section_heading(doc, "DO OBJETO")
        end_str = ", ".join(
            p for p in [
                objeto.get("endereco"), objeto.get("bairro"),
                objeto.get("cidade"), objeto.get("uf"),
            ] if p
        )
        if end_str:
            _p_bold_value(doc, "Imóvel/Objeto", end_str)
        if objeto.get("matricula"):
            _p_bold_value(doc, "Matrícula", objeto["matricula"])
        doc.add_paragraph()

    # Valor total do negocio
    valor_total = cond.get("valor_total", 0)
    if valor_total:
        _p_bold_value(doc, "Valor Total do Negócio", _format_currency(valor_total))

    doc.add_paragraph()

    # Assinatura
    cidade = contrato.get("cidade_assinatura") or ""
    uf = contrato.get("uf_assinatura") or ""
    loc = f"{cidade}/{uf}" if cidade and uf else (cidade or uf or "___________")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{loc}, {sinal_data}.")
    run.font.name = FONT_JURIDICA
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Linha de assinatura do recebedor (Vendedor/Locador)
    recebedor = next((
        p for p in partes
        if any(k in (p.get("qualificacao") or "").lower()
               for k in ["vendedor", "locador", "cedente", "comodante"])
    ), partes[0] if partes else None)

    if recebedor:
        if recebedor.get("tipo") == "pf" and recebedor.get("pf"):
            nome_r = recebedor["pf"].get("nome", "")
        elif recebedor.get("tipo") == "pj" and recebedor.get("pj"):
            nome_r = recebedor["pj"].get("razao_social", "")
        else:
            nome_r = ""
        qual_r = recebedor.get("qualificacao", "RECEBEDOR")
    else:
        nome_r = ""
        qual_r = "RECEBEDOR"

    p = doc.add_paragraph("_" * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    if nome_r:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(nome_r)
        r.font.name = FONT_JURIDICA
        r.font.size = Pt(12)
        r.font.bold = True
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(qual_r.upper() + " — RECEBEDOR DO SINAL")
    r3.font.name = FONT_JURIDICA
    r3.font.size = Pt(11)
    r3.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Entry point principal ─────────────────────────────────────────────────────

def generate_contrato_docx(
    contrato: dict,
    user: dict,
    perfil: Optional[dict] = None,
) -> bytes:
    """Gera DOCX completo do Contrato Imobiliario com layout juridico premium.

    Layout: Times New Roman 12pt | margens 3/2/3/2cm | clausulas em romano
    Retorna bytes do arquivo DOCX.
    """
    if user is None:
        user = {}

    doc = Document()

    # Margens juridicas: esq 3cm, sup 3cm, dir 2cm, inf 2cm
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Define estilo Normal com Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = FONT_JURIDICA
    style.font.size = Pt(12)

    # ── Capa ──────────────────────────────────────────────────────────────────
    _render_cover(doc, contrato, user)

    # ── Qualificacao das Partes ───────────────────────────────────────────────
    _render_qualificacao_partes(doc, contrato)

    # ── Objeto do Contrato ────────────────────────────────────────────────────
    _render_objeto(doc, contrato)

    # ── Condicoes de Pagamento ────────────────────────────────────────────────
    _render_pagamento(doc, contrato)

    # ── Clausulas ─────────────────────────────────────────────────────────────
    _render_clausulas(doc, contrato)

    # ── Corretagem ────────────────────────────────────────────────────────────
    _render_corretor(doc, contrato)

    # ── Foro ──────────────────────────────────────────────────────────────────
    _render_foro(doc, contrato)

    # ── Assinaturas ───────────────────────────────────────────────────────────
    doc.add_page_break()
    _render_assinaturas(doc, contrato, user)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

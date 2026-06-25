# @module services.georef.generators.docx — Versões DOCX editáveis dos documentos.
#
# Reusa o mesmo conteúdo estruturado (textos.py) que o PDF, no padrão python-docx
# do projeto (verde #1B4D1B / dourado #D4A830). Saída em bytes (BytesIO).
import io

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.georef.generators import textos as TX

GREEN = RGBColor(0x0C, 0x33, 0x20)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _p(doc, text, bold=False, italic=False, size=10.5, color=DARK,
       align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for i, bloco in enumerate(str(text or "").split("\n")):
        run = p.add_run(("" if i == 0 else "\n") + bloco)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
    return p


def _titulo(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = text
    _set_cell_bg(cell, "0C3320")
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.runs[0]
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()


def _secao(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = GREEN


def _kv(doc, pares):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in pares:
        row = table.add_row().cells
        rk = row[0].paragraphs[0].add_run(str(k))
        rk.font.bold = True
        rk.font.size = Pt(9)
        rk.font.color.rgb = GREEN
        rv = row[1].paragraphs[0].add_run(str(v))
        rv.font.size = Pt(9)
    doc.add_paragraph()


def _tabela(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        _set_cell_bg(hdr[i], "0C3320")
        run = hdr[i].paragraphs[0].add_run(str(h))
        run.font.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = WHITE
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in rows:
        cells = table.add_row().cells
        for i, c in enumerate(r):
            run = cells[i].paragraphs[0].add_run("" if c is None else str(c))
            run.font.size = Pt(7.5)
    doc.add_paragraph()


def _assinaturas(doc, assinaturas):
    doc.add_paragraph()
    for nome, papel in assinaturas:
        doc.add_paragraph()
        p = doc.add_paragraph("____________________________________________")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pn = doc.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = pn.add_run(str(nome))
        rn.font.bold = True
        rn.font.size = Pt(9.5)
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp = pp.add_run(str(papel))
        rp.font.size = Pt(9)


def _save(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
def docx_requerimento(projeto) -> bytes:
    d = TX.render_requerimento(projeto)
    doc = Document()
    _titulo(doc, d["titulo"])
    _p(doc, d["destinatario"], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(doc, d["corpo"])
    for item in d["documentos"]:
        _p(doc, item, space_after=2)
    _p(doc, d["fecho"])
    _p(doc, d["rt_linha"], size=8.5, color=RGBColor(0x55, 0x55, 0x55))
    _p(doc, "Nestes termos, pede deferimento.")
    _p(doc, d["data"] + ".", align=WD_ALIGN_PARAGRAPH.CENTER)
    _assinaturas(doc, [(d["assinatura"][0], d["assinatura"][1])])
    return _save(doc)


def docx_memorial(projeto) -> bytes:
    d = TX.render_memorial(projeto)
    doc = Document()
    _titulo(doc, d["titulo"])
    _secao(doc, "IDENTIFICAÇÃO")
    _kv(doc, d["cabecalho"])
    _secao(doc, "DESCRIÇÃO DA PARCELA (VÉRTICES)")
    _tabela(doc, d["tabela_header"], d["tabela"])
    _p(doc, d["data"] + ".", align=WD_ALIGN_PARAGRAPH.CENTER)
    _assinaturas(doc, [(d["rt_assinatura"][0], " — ".join(d["rt_assinatura"][1:]))])
    return _save(doc)


def docx_drl(projeto, conf) -> bytes:
    d = TX.render_drl(projeto, conf)
    doc = Document()
    _titulo(doc, d["titulo"])
    for linha in d["cabecalho"]:
        _p(doc, linha, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    _p(doc, d["corpo"])
    _secao(doc, "SEGMENTOS DIVISÓRIOS (SIRGAS 2000)")
    if d["tabela"]:
        _tabela(doc, d["tabela_header"], d["tabela"])
    else:
        _p(doc, "Sem segmentos vinculados a este confrontante.", size=8.5)
    _p(doc, d["data"] + ".", align=WD_ALIGN_PARAGRAPH.CENTER)
    _assinaturas(doc, d["assinaturas"])
    return _save(doc)


def docx_laudo(projeto) -> bytes:
    d = TX.render_laudo_tecnico(projeto)
    doc = Document()
    _titulo(doc, d["titulo"])
    _secao(doc, "1. IDENTIFICAÇÃO")
    _p(doc, d["identificacao"])
    _secao(doc, "2. OBJETO")
    _p(doc, d["objeto"])
    _secao(doc, "3. JUSTIFICATIVA TÉCNICO-JURÍDICA")
    _p(doc, d["justificativa"])
    _secao(doc, "4. METODOLOGIA E PARÂMETROS TÉCNICOS DE AGRIMENSURA")
    _p(doc, d["metodologia"])
    _secao(doc, "5. RESULTADO — POLIGONAL")
    _tabela(doc, d["resultado_header"], d["resultado_tabela"])
    _secao(doc, "6. CADEIA DOMINIAL")
    if d["cadeia_tabela"]:
        _tabela(doc, d["cadeia_header"], d["cadeia_tabela"])
    else:
        _p(doc, d["cadeia_nota"])
    _secao(doc, "7. CONFRONTAÇÕES")
    _tabela(doc, d["confrontacoes_header"], d["confrontacoes_tabela"])
    _secao(doc, "8. CONCLUSÃO")
    _p(doc, d["conclusao"])
    _p(doc, d["data"] + ".", align=WD_ALIGN_PARAGRAPH.CENTER)
    _assinaturas(doc, [(d["rt_assinatura"][0], " — ".join(d["rt_assinatura"][1:]))])
    return _save(doc)


def gerar_docx(tipo, projeto, conf=None) -> bytes:
    if tipo == "requerimento":
        return docx_requerimento(projeto)
    if tipo == "memorial":
        return docx_memorial(projeto)
    if tipo == "laudo_tecnico":
        return docx_laudo(projeto)
    if tipo == "drl":
        if conf is None:
            raise ValueError("DRL requer o confrontante (conf).")
        return docx_drl(projeto, conf)
    raise ValueError(f"Tipo de documento desconhecido: {tipo}")
